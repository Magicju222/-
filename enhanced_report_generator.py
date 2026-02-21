"""
增强版报告生成器
直接从 Agent 分析结果中提取完整数据生成报告
确保 Word 和 Excel 报告中的数据完整且以表格形式呈现
"""

import io
import re
import json
import pandas as pd
import numpy as np
from datetime import datetime
from typing import Dict, List, Any, Optional, Tuple
from dataclasses import dataclass, field

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.utils.dataframe import dataframe_to_rows


@dataclass
class ExtractedTable:
    """提取的表格数据"""
    title: str
    headers: List[str]
    rows: List[List[Any]]
    source: str = ""  # 数据来源步骤


@dataclass
class AnalysisSection:
    """分析章节"""
    title: str
    content: str
    tables: List[ExtractedTable] = field(default_factory=list)


class DataExtractor:
    """数据提取器 - 从 Agent 结果中提取结构化数据"""
    
    def __init__(self, agent_result):
        self.agent_result = agent_result
        self.all_tables: List[ExtractedTable] = []
        self.sections: List[AnalysisSection] = []
        
    def extract_all_data(self):
        """提取所有数据"""
        self._extract_from_steps()
        self._extract_from_final_report()
        return self
    
    def _extract_from_steps(self):
        """从分析步骤中提取数据"""
        if not hasattr(self.agent_result, 'steps') or not self.agent_result.steps:
            return
        
        for step in self.agent_result.steps:
            observation = getattr(step, 'observation', '')
            if not observation:
                continue
            
            # 提取表格
            tables = self._parse_tables_from_text(observation)
            for table in tables:
                table.source = f"步骤 {getattr(step, 'step_number', 0)}"
                self.all_tables.append(table)
            
            # 创建章节
            if tables or len(observation.strip()) > 50:
                section = AnalysisSection(
                    title=f"分析步骤 {getattr(step, 'step_number', 0)}",
                    content=self._extract_text_content(observation),
                    tables=tables
                )
                self.sections.append(section)
    
    def _extract_from_final_report(self):
        """从最终报告中提取数据"""
        if not hasattr(self.agent_result, 'final_report') or not self.agent_result.final_report:
            return
        
        final_report = self.agent_result.final_report
        
        # 提取表格
        tables = self._parse_tables_from_text(final_report)
        for table in tables:
            table.source = "最终报告"
            self.all_tables.append(table)
        
        # 创建章节
        if tables:
            section = AnalysisSection(
                title="综合分析结果",
                content=self._extract_text_content(final_report),
                tables=tables
            )
            self.sections.append(section)
    
    def _parse_tables_from_text(self, text: str) -> List[ExtractedTable]:
        """从文本中解析表格"""
        tables = []
        lines = text.strip().split('\n')
        
        current_table_lines = []
        table_title = ""
        in_table = False
        
        for i, line in enumerate(lines):
            stripped = line.strip()
            
            # 检测表格标题（在表格之前的行）
            if not in_table and stripped and not self._is_table_line(stripped):
                if i + 1 < len(lines) and self._is_table_line(lines[i + 1].strip()):
                    table_title = stripped.replace('**', '').replace('#', '').strip()
            
            # 检测表格行
            if self._is_table_line(stripped):
                if not in_table:
                    in_table = True
                    current_table_lines = []
                current_table_lines.append(stripped)
            else:
                if in_table and current_table_lines:
                    # 解析当前表格
                    table = self._parse_table_lines(current_table_lines, table_title)
                    if table:
                        tables.append(table)
                    current_table_lines = []
                    in_table = False
                    table_title = ""
        
        # 处理最后一个表格
        if in_table and current_table_lines:
            table = self._parse_table_lines(current_table_lines, table_title)
            if table:
                tables.append(table)
        
        return tables
    
    def _is_table_line(self, line: str) -> bool:
        """判断是否为表格行"""
        # Markdown 表格格式
        if '|' in line:
            parts = [p.strip() for p in line.split('|')]
            parts = [p for p in parts if p]
            return len(parts) >= 2
        # 制表符分隔
        if '\t' in line:
            parts = line.split('\t')
            return len(parts) >= 2
        return False
    
    def _parse_table_lines(self, table_lines: List[str], title: str = "") -> Optional[ExtractedTable]:
        """解析表格行"""
        if len(table_lines) < 2:
            return None
        
        headers = []
        rows = []
        
        for i, line in enumerate(table_lines):
            # 分割单元格
            if '|' in line:
                cells = [cell.strip() for cell in line.split('|')]
                cells = [c for c in cells if c or c == '']
            else:
                cells = [cell.strip() for cell in line.split('\t')]
            
            if not cells or all(c == '' for c in cells):
                continue
            
            # 跳过 Markdown 分隔符行
            if all(c.replace('-', '').replace(':', '').strip() == '' for c in cells):
                continue
            
            if i == 0 or not headers:
                headers = cells
            else:
                # 确保单元格数与表头一致
                if len(cells) < len(headers):
                    cells.extend([''] * (len(headers) - len(cells)))
                elif len(cells) > len(headers):
                    cells = cells[:len(headers)]
                rows.append(cells)
        
        if headers and rows:
            return ExtractedTable(
                title=title or "数据表格",
                headers=headers,
                rows=rows
            )
        
        return None
    
    def _extract_text_content(self, text: str) -> str:
        """提取非表格文本内容"""
        lines = text.strip().split('\n')
        text_lines = []
        in_table = False
        
        for line in lines:
            stripped = line.strip()
            
            if self._is_table_line(stripped):
                in_table = True
                continue
            else:
                if in_table:
                    in_table = False
                if stripped and not stripped.startswith('---'):
                    text_lines.append(stripped)
        
        return '\n'.join(text_lines)
    
    def get_summary_statistics(self) -> Dict[str, Any]:
        """获取汇总统计信息"""
        return {
            'total_tables': len(self.all_tables),
            'total_rows': sum(len(t.rows) for t in self.all_tables),
            'total_sections': len(self.sections),
            'tables_by_section': [
                {
                    'section': s.title,
                    'table_count': len(s.tables),
                    'row_count': sum(len(t.rows) for t in s.tables)
                }
                for s in self.sections
            ]
        }


class EnhancedWordReportGenerator:
    """增强版 Word 报告生成器"""
    
    def __init__(self, extractor: DataExtractor):
        self.extractor = extractor
        self.doc = Document()
        self._setup_document()
    
    def _setup_document(self):
        """设置文档样式"""
        # 设置默认字体
        style = self.doc.styles['Normal']
        style.font.name = 'Microsoft YaHei'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        style.font.size = Pt(10.5)
    
    def _set_cell_font(self, cell, font_name='Microsoft YaHei', font_size=10, bold=False, color=None):
        """设置单元格字体"""
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = font_name
                run.font.size = Pt(font_size)
                run.font.bold = bold
                run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                if color:
                    run.font.color.rgb = color
    
    def generate(self, title: str = "数据分析报告", sheet_name: str = "") -> io.BytesIO:
        """生成 Word 报告"""
        # 标题
        title_heading = self.doc.add_heading(title, 0)
        for run in title_heading.runs:
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            run.font.size = Pt(18)
            run.font.bold = True
        title_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 元信息
        if sheet_name:
            p = self.doc.add_paragraph()
            run = p.add_run(f"分析对象: {sheet_name}")
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        
        p = self.doc.add_paragraph()
        run = p.add_run(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        
        self.doc.add_paragraph()
        
        # 数据概览
        stats = self.extractor.get_summary_statistics()
        self._add_summary_section(stats)
        
        # 详细数据表格
        self._add_data_tables()
        
        # 保存
        doc_io = io.BytesIO()
        self.doc.save(doc_io)
        doc_io.seek(0)
        return doc_io
    
    def _add_summary_section(self, stats: Dict):
        """添加汇总章节"""
        heading = self.doc.add_heading('一、数据概览', level=1)
        for run in heading.runs:
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            run.font.size = Pt(14)
            run.font.bold = True
        
        # 创建汇总表格
        overview_data = [
            ['统计项目', '数值', '说明'],
            ['数据表格数', stats['total_tables'], '报告中包含的表格数量'],
            ['数据总行数', stats['total_rows'], '所有表格的数据行总数'],
            ['分析章节数', stats['total_sections'], '报告的分析章节数量'],
        ]
        
        table = self.doc.add_table(rows=len(overview_data), cols=3)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        for i, row_data in enumerate(overview_data):
            row = table.rows[i]
            for j, value in enumerate(row_data):
                cell = row.cells[j]
                cell.text = str(value)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # 表头样式
                if i == 0:
                    self._set_cell_font(cell, bold=True, color=RGBColor(255, 255, 255))
                    shading_elm = parse_xml(r'<w:shd {} w:fill="4472C4"/>'.format(nsdecls('w')))
                    cell._tc.get_or_add_tcPr().append(shading_elm)
                else:
                    self._set_cell_font(cell)
        
        self.doc.add_paragraph()
        
        # 各章节统计
        if stats['tables_by_section']:
            p = self.doc.add_paragraph()
            run = p.add_run("各章节数据统计：")
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            run.font.bold = True
            
            for section_stat in stats['tables_by_section']:
                p = self.doc.add_paragraph(
                    f"• {section_stat['section']}: {section_stat['table_count']} 个表格，"
                    f"共 {section_stat['row_count']} 行数据",
                    style='List Bullet'
                )
                for run in p.runs:
                    run.font.name = 'Microsoft YaHei'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        
        self.doc.add_page_break()
    
    def _add_data_tables(self):
        """添加所有数据表格"""
        heading = self.doc.add_heading('二、详细数据分析', level=1)
        for run in heading.runs:
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            run.font.size = Pt(14)
            run.font.bold = True
        
        if not self.extractor.all_tables:
            p = self.doc.add_paragraph("未找到数据表格。")
            for run in p.runs:
                run.font.name = 'Microsoft YaHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            return
        
        # 按章节分组显示表格
        for section in self.extractor.sections:
            if not section.tables:
                continue
            
            # 章节标题
            section_heading = self.doc.add_heading(section.title, level=2)
            for run in section_heading.runs:
                run.font.name = 'Microsoft YaHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                run.font.size = Pt(12)
                run.font.bold = True
            
            # 章节内容
            if section.content:
                p = self.doc.add_paragraph(section.content)
                for run in p.runs:
                    run.font.name = 'Microsoft YaHei'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                    run.font.size = Pt(10)
                self.doc.add_paragraph()
            
            # 显示该章节的所有表格
            for table_idx, table_data in enumerate(section.tables, 1):
                self._add_single_table(table_data, table_idx)
            
            self.doc.add_paragraph()
    
    def _add_single_table(self, table_data: ExtractedTable, table_idx: int):
        """添加单个表格"""
        # 表格标题
        p = self.doc.add_paragraph()
        run = p.add_run(f"表 {table_idx}: {table_data.title}")
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        run.font.size = Pt(11)
        run.font.bold = True
        
        # 创建表格
        if not table_data.headers or not table_data.rows:
            p = self.doc.add_paragraph("(表格数据为空)")
            for run in p.runs:
                run.font.name = 'Microsoft YaHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            return
        
        table = self.doc.add_table(rows=1 + len(table_data.rows), cols=len(table_data.headers))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 表头
        hdr_cells = table.rows[0].cells
        for j, header in enumerate(table_data.headers):
            hdr_cells[j].text = str(header)
            hdr_cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._set_cell_font(hdr_cells[j], bold=True, color=RGBColor(255, 255, 255))
            # 深蓝色背景
            shading_elm = parse_xml(r'<w:shd {} w:fill="1F4E79"/>'.format(nsdecls('w')))
            hdr_cells[j]._tc.get_or_add_tcPr().append(shading_elm)
        
        # 数据行
        for row_idx, row_data in enumerate(table_data.rows, 1):
            row_cells = table.rows[row_idx].cells
            for j, cell_value in enumerate(row_data):
                if j < len(row_cells):
                    row_cells[j].text = str(cell_value)
                    row_cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    self._set_cell_font(row_cells[j], font_size=9)
                    
                    # 交替行背景色
                    if row_idx % 2 == 0:
                        shading_elm = parse_xml(r'<w:shd {} w:fill="F2F2F2"/>'.format(nsdecls('w')))
                        row_cells[j]._tc.get_or_add_tcPr().append(shading_elm)
        
        # 表格说明
        p = self.doc.add_paragraph()
        run = p.add_run(f"*本表包含 {len(table_data.rows)} 行数据，{len(table_data.headers)} 个字段*")
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        run.font.size = Pt(9)
        run.font.italic = True
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        self.doc.add_paragraph()


class EnhancedExcelReportGenerator:
    """增强版 Excel 报告生成器"""
    
    def __init__(self, extractor: DataExtractor):
        self.extractor = extractor
        self.wb = Workbook()
        if 'Sheet' in self.wb.sheetnames:
            self.wb.remove(self.wb['Sheet'])
    
    def _format_header(self, cell):
        """格式化表头"""
        cell.font = Font(name='微软雅黑', size=11, bold=True, color='FFFFFF')
        cell.fill = PatternFill(start_color='4472C4', end_color='4472C4', fill_type='solid')
        cell.alignment = Alignment(horizontal='center', vertical='center')
        cell.border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
    
    def _format_data_cell(self, cell, bold: bool = False):
        """格式化数据单元格"""
        cell.font = Font(name='微软雅黑', size=10, bold=bold)
        cell.alignment = Alignment(horizontal='center', vertical='center')
        cell.border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
    
    def generate(self, title: str = "数据分析报告") -> io.BytesIO:
        """生成 Excel 报告"""
        # Sheet 1: 概览
        self._create_overview_sheet(title)
        
        # Sheet 2+: 数据表格
        self._create_data_sheets()
        
        # 保存
        output = io.BytesIO()
        self.wb.save(output)
        output.seek(0)
        return output
    
    def _create_overview_sheet(self, title: str):
        """创建概览工作表"""
        ws = self.wb.create_sheet("报告概览")
        current_row = 1
        
        # 标题
        cell = ws.cell(row=current_row, column=1, value=title)
        cell.font = Font(name='微软雅黑', size=16, bold=True)
        ws.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=5)
        current_row += 2
        
        # 生成时间
        cell = ws.cell(row=current_row, column=1, value=f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        cell.font = Font(name='微软雅黑', size=10)
        current_row += 2
        
        # 统计信息
        stats = self.extractor.get_summary_statistics()
        
        cell = ws.cell(row=current_row, column=1, value="数据概览")
        cell.font = Font(name='微软雅黑', size=12, bold=True)
        current_row += 1
        
        overview_data = [
            ['统计项目', '数值', '说明'],
            ['数据表格数', stats['total_tables'], '报告中包含的表格数量'],
            ['数据总行数', stats['total_rows'], '所有表格的数据行总数'],
            ['分析章节数', stats['total_sections'], '报告的分析章节数量'],
        ]
        
        for row_idx, row_data in enumerate(overview_data, current_row):
            for col_idx, value in enumerate(row_data, 1):
                cell = ws.cell(row=row_idx, column=col_idx, value=value)
                if row_idx == current_row:
                    self._format_header(cell)
                else:
                    self._format_data_cell(cell)
        
        current_row += len(overview_data) + 2
        
        # 各章节统计
        if stats['tables_by_section']:
            cell = ws.cell(row=current_row, column=1, value="各章节数据统计")
            cell.font = Font(name='微软雅黑', size=12, bold=True)
            current_row += 1
            
            section_headers = ['章节名称', '表格数量', '数据行数']
            for col_idx, header in enumerate(section_headers, 1):
                cell = ws.cell(row=current_row, column=col_idx, value=header)
                self._format_header(cell)
            current_row += 1
            
            for section_stat in stats['tables_by_section']:
                row_data = [
                    section_stat['section'],
                    section_stat['table_count'],
                    section_stat['row_count']
                ]
                for col_idx, value in enumerate(row_data, 1):
                    cell = ws.cell(row=current_row, column=col_idx, value=value)
                    self._format_data_cell(cell)
                current_row += 1
        
        # 调整列宽
        ws.column_dimensions['A'].width = 20
        ws.column_dimensions['B'].width = 15
        ws.column_dimensions['C'].width = 40
    
    def _create_data_sheets(self):
        """创建数据工作表"""
        if not self.extractor.all_tables:
            ws = self.wb.create_sheet("数据详情")
            ws.cell(row=1, column=1, value="未找到数据表格")
            return
        
        # 创建一个汇总表，包含所有数据
        ws_all = self.wb.create_sheet("所有数据")
        current_row = 1
        
        for table_idx, table_data in enumerate(self.extractor.all_tables, 1):
            # 表格标题
            cell = ws_all.cell(row=current_row, column=1, value=f"表 {table_idx}: {table_data.title}")
            cell.font = Font(name='微软雅黑', size=12, bold=True)
            ws_all.merge_cells(start_row=current_row, start_column=1, 
                              end_row=current_row, end_column=len(table_data.headers))
            current_row += 1
            
            # 表头
            for col_idx, header in enumerate(table_data.headers, 1):
                cell = ws_all.cell(row=current_row, column=col_idx, value=header)
                self._format_header(cell)
            current_row += 1
            
            # 数据行
            for row_data in table_data.rows:
                for col_idx, value in enumerate(row_data, 1):
                    cell = ws_all.cell(row=current_row, column=col_idx, value=value)
                    self._format_data_cell(cell)
                current_row += 1
            
            current_row += 2  # 表格间距
        
        # 为每个章节创建单独的工作表
        for section in self.extractor.sections:
            if not section.tables:
                continue
            
            # 工作表名称（限制长度）
            sheet_name = section.title[:31]  # Excel 工作表名称最大 31 字符
            ws = self.wb.create_sheet(sheet_name)
            current_row = 1
            
            # 章节标题
            cell = ws.cell(row=current_row, column=1, value=section.title)
            cell.font = Font(name='微软雅黑', size=14, bold=True)
            ws.merge_cells(start_row=current_row, start_column=1, 
                          end_row=current_row, end_column=10)
            current_row += 2
            
            # 章节内容
            if section.content:
                cell = ws.cell(row=current_row, column=1, value=section.content)
                cell.font = Font(name='微软雅黑', size=10)
                ws.merge_cells(start_row=current_row, start_column=1, 
                              end_row=current_row, end_column=10)
                current_row += 2
            
            # 表格
            for table_data in section.tables:
                # 表格标题
                cell = ws.cell(row=current_row, column=1, value=table_data.title)
                cell.font = Font(name='微软雅黑', size=11, bold=True)
                ws.merge_cells(start_row=current_row, start_column=1, 
                              end_row=current_row, end_column=len(table_data.headers))
                current_row += 1
                
                # 表头
                for col_idx, header in enumerate(table_data.headers, 1):
                    cell = ws.cell(row=current_row, column=col_idx, value=header)
                    self._format_header(cell)
                current_row += 1
                
                # 数据行
                for row_data in table_data.rows:
                    for col_idx, value in enumerate(row_data, 1):
                        cell = ws.cell(row=current_row, column=col_idx, value=value)
                        self._format_data_cell(cell)
                    current_row += 1
                
                current_row += 1
            
            # 调整列宽
            for col in range(1, 11):
                ws.column_dimensions[chr(64 + col) if col <= 26 else 'A'].width = 15


def generate_enhanced_reports(agent_result, sheet_name: str = "") -> Tuple[str, io.BytesIO, io.BytesIO]:
    """
    生成增强版报告
    
    Returns:
        (markdown_report, word_report, excel_report)
    """
    # 提取数据
    extractor = DataExtractor(agent_result)
    extractor.extract_all_data()
    
    # 生成 Markdown 报告
    md_generator = EnhancedMarkdownReportGenerator(extractor)
    md_report = md_generator.generate(sheet_name)
    
    # 生成 Word 报告
    word_generator = EnhancedWordReportGenerator(extractor)
    word_report = word_generator.generate(sheet_name=sheet_name)
    
    # 生成 Excel 报告
    excel_generator = EnhancedExcelReportGenerator(extractor)
    excel_report = excel_generator.generate()
    
    return md_report, word_report, excel_report


class EnhancedMarkdownReportGenerator:
    """增强版 Markdown 报告生成器"""
    
    def __init__(self, extractor: DataExtractor):
        self.extractor = extractor
    
    def generate(self, sheet_name: str = "") -> str:
        """生成 Markdown 报告"""
        lines = []
        stats = self.extractor.get_summary_statistics()
        
        # 标题
        lines.append("# 数据分析报告")
        lines.append("")
        if sheet_name:
            lines.append(f"**分析对象**: {sheet_name}")
        lines.append(f"**生成时间**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        lines.append("")
        lines.append("---")
        lines.append("")
        
        # 数据概览
        lines.append("## 一、数据概览")
        lines.append("")
        lines.append(f"- **数据表格数**: {stats['total_tables']}")
        lines.append(f"- **数据总行数**: {stats['total_rows']}")
        lines.append(f"- **分析章节数**: {stats['total_sections']}")
        lines.append("")
        
        if stats['tables_by_section']:
            lines.append("### 各章节数据统计")
            lines.append("")
            for section_stat in stats['tables_by_section']:
                lines.append(f"- **{section_stat['section']}**: {section_stat['table_count']} 个表格，"
                           f"共 {section_stat['row_count']} 行数据")
            lines.append("")
        
        lines.append("---")
        lines.append("")
        
        # 详细数据
        lines.append("## 二、详细数据分析")
        lines.append("")
        
        if not self.extractor.sections:
            lines.append("未找到分析数据。")
        else:
            for section in self.extractor.sections:
                if not section.tables:
                    continue
                
                lines.append(f"### {section.title}")
                lines.append("")
                
                if section.content:
                    lines.append(section.content)
                    lines.append("")
                
                for table_idx, table in enumerate(section.tables, 1):
                    lines.append(f"#### 表 {table_idx}: {table.title}")
                    lines.append("")
                    
                    # Markdown 表格
                    if table.headers and table.rows:
                        lines.append("| " + " | ".join(table.headers) + " |")
                        lines.append("|" + "|".join(["---"] * len(table.headers)) + "|")
                        for row in table.rows:
                            lines.append("| " + " | ".join([str(c) for c in row]) + " |")
                        lines.append("")
                        lines.append(f"*本表包含 {len(table.rows)} 行数据*")
                        lines.append("")
                
                lines.append("---")
                lines.append("")
        
        return "\n".join(lines)


# 便捷函数
def generate_enhanced_word_report(agent_result, sheet_name: str = "") -> io.BytesIO:
    """生成增强版 Word 报告"""
    extractor = DataExtractor(agent_result)
    extractor.extract_all_data()
    generator = EnhancedWordReportGenerator(extractor)
    return generator.generate(sheet_name=sheet_name)


def generate_enhanced_excel_report(agent_result, sheet_name: str = "") -> io.BytesIO:
    """生成增强版 Excel 报告"""
    extractor = DataExtractor(agent_result)
    extractor.extract_all_data()
    generator = EnhancedExcelReportGenerator(extractor)
    return generator.generate()

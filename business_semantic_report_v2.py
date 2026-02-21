"""
业务语义分析报告生成器 V2
基于数据自动识别业务指标、维度、趋势，生成专业的业务分析报告
改进版本：修复维度分析空白、洞察空白、表格样式、添加详细说明
"""

import io
import json
import re
from typing import Any, Dict, List, Optional, Union, Tuple
from datetime import datetime
from dataclasses import dataclass, field

import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from llm_client import LLMClient, LLMConfig


@dataclass
class BusinessMetric:
    """业务指标"""
    name: str
    value: float
    unit: str = ""
    description: str = ""
    aggregation_type: str = ""  # sum, avg, max, min, std, count
    format_str: str = "{:.2f}"


@dataclass
class DimensionAnalysis:
    """维度分析结果"""
    dimension_name: str
    metric_name: str
    data: pd.DataFrame  # 包含维度值和对应指标值
    summary: str = ""  # 分析总结
    top_values: List[Dict] = field(default_factory=list)
    bottom_values: List[Dict] = field(default_factory=list)


@dataclass
class TrendAnalysis:
    """趋势分析结果"""
    time_field: str
    metric_name: str
    data: pd.DataFrame  # 包含时间和指标值
    summary: str = ""  # 趋势总结
    trend_direction: str = ""  # up, down, stable
    growth_rate: float = 0.0
    peak_point: Optional[Dict] = None
    valley_point: Optional[Dict] = None


@dataclass
class KeyInsight:
    """关键洞察"""
    title: str
    description: str
    insight_type: str  # anomaly, opportunity, pattern, risk
    supporting_data: Dict[str, Any]
    confidence: str = "中"
    priority: str = "中"


@dataclass
class Recommendation:
    """业务建议"""
    title: str
    description: str
    recommendation_type: str  # improvement, opportunity, optimization
    expected_impact: str = ""
    implementation_difficulty: str = "中"
    supporting_insights: List[str] = field(default_factory=list)


class BusinessSemanticReportGeneratorV2:
    """
    业务语义分析报告生成器 V2
    自动识别业务指标、维度、趋势，生成专业的业务分析报告
    """
    
    def __init__(self, df: Optional[pd.DataFrame] = None, llm_client: Optional[LLMClient] = None):
        self.llm_client = llm_client
        self.df: Optional[pd.DataFrame] = None
        self.numeric_columns: List[str] = []
        self.categorical_columns: List[str] = []
        self.datetime_columns: List[str] = []
        self.business_metrics: List[BusinessMetric] = []
        self.dimension_analyses: List[DimensionAnalysis] = []
        self.trend_analyses: List[TrendAnalysis] = []
        self.key_insights: List[KeyInsight] = []
        self.recommendations: List[Recommendation] = []
        
        # 如果提供了DataFrame，立即检测列类型
        if df is not None:
            self._detect_column_types(df)
        
    def _detect_column_types(self, df: pd.DataFrame) -> None:
        """检测列类型"""
        self.df = df.copy()
        self.numeric_columns = []
        self.categorical_columns = []
        self.datetime_columns = []
        
        for col in df.columns:
            # 跳过明显是ID的列
            if any(keyword in col.lower() for keyword in ['id', '编号', '序号', 'code']):
                continue
                
            if pd.api.types.is_datetime64_any_dtype(df[col]):
                self.datetime_columns.append(col)
            elif pd.api.types.is_numeric_dtype(df[col]):
                # 检查是否为分类编码（数值范围很小）
                unique_count = df[col].nunique()
                total_count = len(df)
                if unique_count <= 10 and unique_count / total_count < 0.1:
                    self.categorical_columns.append(col)
                else:
                    self.numeric_columns.append(col)
            else:
                # 尝试转换为日期
                try:
                    pd.to_datetime(df[col], errors='raise')
                    self.datetime_columns.append(col)
                except:
                    # 检查是否为分类变量
                    unique_count = df[col].nunique()
                    total_count = len(df)
                    if unique_count <= 50 or unique_count / total_count < 0.3:
                        self.categorical_columns.append(col)
        
        print(f"[检测列类型] 数值列: {self.numeric_columns}, 分类列: {self.categorical_columns}, 时间列: {self.datetime_columns}")
    
    def _get_metric_unit(self, column_name: str) -> str:
        """根据列名推断单位"""
        name_lower = column_name.lower()
        
        # 金额相关
        if any(kw in name_lower for kw in ['金额', '价格', '费用', '成本', '收入', '销售额', '营收', 'price', 'amount', 'cost', 'revenue', 'sales']):
            return "元"
        # 数量相关
        elif any(kw in name_lower for kw in ['数量', '销量', 'count', 'quantity', 'volume', 'qty']):
            return "个/件"
        # 百分比相关
        elif any(kw in name_lower for kw in ['率', '比例', '占比', 'percent', 'rate', 'ratio']):
            return "%"
        # 时间相关
        elif any(kw in name_lower for kw in ['时间', '时长', '周期', 'time', 'duration']):
            return "小时"
        # 人数相关
        elif any(kw in name_lower for kw in ['人数', '员工', '客户', 'people', 'staff', 'customer']):
            return "人"
        else:
            return ""
    
    def _get_metric_description(self, column_name: str, aggregation_type: str) -> str:
        """生成指标描述"""
        descriptions = {
            'sum': f"{column_name}的总和",
            'avg': f"{column_name}的平均值",
            'max': f"{column_name}的最大值",
            'min': f"{column_name}的最小值",
            'std': f"{column_name}的标准差",
            'count': f"{column_name}的计数"
        }
        return descriptions.get(aggregation_type, f"{column_name}的{aggregation_type}")
    
    def _analyze_business_metrics(self) -> List[BusinessMetric]:
        """分析业务指标"""
        metrics = []
        
        if self.df is None or self.df.empty:
            return metrics
        
        for col in self.numeric_columns:
            series = self.df[col].dropna()
            
            if series.empty:
                continue
            
            unit = self._get_metric_unit(col)
            
            # 计算各种聚合指标
            aggregations = [
                ('sum', series.sum(), "{:.2f}"),
                ('avg', series.mean(), "{:.2f}"),
                ('max', series.max(), "{:.2f}"),
                ('min', series.min(), "{:.2f}"),
            ]
            
            # 根据列名特征决定主要展示哪些指标
            name_lower = col.lower()
            
            # 金额类指标：优先展示sum和avg
            if any(kw in name_lower for kw in ['金额', '价格', '费用', '成本', '收入', '销售额', '营收', 'price', 'amount', 'cost', 'revenue', 'sales']):
                priority_aggs = ['sum', 'avg', 'max', 'min']
            # 其他：展示全部
            else:
                priority_aggs = ['sum', 'avg', 'max', 'min']
            
            for agg_type, value, fmt in aggregations:
                if agg_type in priority_aggs:
                    metric = BusinessMetric(
                        name=f"{col} ({self._get_agg_name(agg_type)})",
                        value=float(value),
                        unit=unit,
                        description=self._get_metric_description(col, agg_type),
                        aggregation_type=agg_type,
                        format_str=fmt
                    )
                    metrics.append(metric)
        
        self.business_metrics = metrics
        print(f"[业务指标分析] 识别到 {len(metrics)} 个指标")
        return metrics
    
    def _get_agg_name(self, agg_type: str) -> str:
        """获取聚合类型中文名"""
        names = {
            'sum': '汇总',
            'avg': '平均',
            'max': '最大',
            'min': '最小',
            'std': '标准差',
            'count': '计数'
        }
        return names.get(agg_type, agg_type)
    
    def _analyze_by_dimensions(self) -> List[DimensionAnalysis]:
        """分析维度 - 修复空白问题"""
        analyses = []
        
        if self.df is None or self.df.empty:
            print("[维度分析] 数据为空")
            return analyses
        
        if not self.categorical_columns:
            print("[维度分析] 没有分类列")
            return analyses
        
        if not self.numeric_columns:
            print("[维度分析] 没有数值列")
            return analyses
        
        print(f"[维度分析] 开始分析，分类列: {self.categorical_columns}, 数值列: {self.numeric_columns}")
        
        # 对每个分类维度进行分析
        for dim_col in self.categorical_columns:
            # 跳过唯一值过多的维度
            unique_count = self.df[dim_col].nunique()
            if unique_count > 100:
                print(f"[维度分析] 跳过 {dim_col}，唯一值过多: {unique_count}")
                continue
            
            # 对每个数值指标进行分析
            for metric_col in self.numeric_columns[:3]:  # 限制指标数量
                try:
                    print(f"[维度分析] 分析 {dim_col} x {metric_col}")
                    
                    # 按维度分组计算指标（显式设置 observed=False 避免警告）
                    grouped = self.df.groupby(dim_col, observed=False)[metric_col].agg([
                        ('汇总', 'sum'),
                        ('平均', 'mean'),
                        ('计数', 'count')
                    ]).reset_index()
                    
                    grouped.columns = [dim_col, '汇总', '平均', '计数']
                    
                    # 按汇总值排序
                    grouped = grouped.sort_values('汇总', ascending=False)
                    
                    # 获取Top和Bottom
                    top_values = grouped.head(3).to_dict('records')
                    bottom_values = grouped.tail(3).to_dict('records')
                    
                    # 生成分析总结
                    total_sum = grouped['汇总'].sum()
                    top1 = top_values[0] if top_values else None
                    if top1:
                        top1_value = top1.get('汇总', 0)
                        top1_name = top1.get(dim_col, '未知')
                        top1_ratio = (top1_value / total_sum * 100) if total_sum > 0 else 0
                        summary = f"{top1_name}表现最优，占总汇总的{top1_ratio:.1f}%（{top1_value:.2f}）"
                    else:
                        summary = f"{dim_col}维度分析完成"
                    
                    analysis = DimensionAnalysis(
                        dimension_name=dim_col,
                        metric_name=metric_col,
                        data=grouped,
                        summary=summary,
                        top_values=top_values,
                        bottom_values=bottom_values
                    )
                    analyses.append(analysis)
                    print(f"[维度分析] 完成 {dim_col} x {metric_col}，数据行数: {len(grouped)}")
                    
                except Exception as e:
                    print(f"[维度分析] 错误 {dim_col} x {metric_col}: {e}")
                    continue
        
        self.dimension_analyses = analyses
        print(f"[维度分析] 总共完成 {len(analyses)} 个维度分析")
        return analyses
    
    def _extract_key_insights(self) -> List[KeyInsight]:
        """提取关键洞察 - 确保不为空"""
        insights = []
        
        print(f"[关键洞察] 开始提取，维度分析数: {len(self.dimension_analyses)}")
        
        # 1. 从维度分析中提取洞察
        for analysis in self.dimension_analyses:
            if analysis.top_values:
                top = analysis.top_values[0]
                dim_name = analysis.dimension_name
                metric_name = analysis.metric_name
                
                # 计算占比
                total = sum(v.get('汇总', 0) for v in analysis.data.to_dict('records'))
                top_value = top.get('汇总', 0)
                ratio = (top_value / total * 100) if total > 0 else 0
                
                insight = KeyInsight(
                    title=f"{dim_name}维度：{top.get(dim_name, '未知')}领先",
                    description=f"在{metric_name}指标上，{top.get(dim_name, '未知')}表现最优，"
                               f"数值为{top_value:.2f}，占总汇总的{ratio:.1f}%。"
                               f"这表明该类别在业务中占据主导地位。",
                    insight_type="opportunity",
                    supporting_data={
                        'dimension': dim_name,
                        'metric': metric_name,
                        'top_category': top.get(dim_name, '未知'),
                        'top_value': top_value,
                        'ratio_percent': round(ratio, 2),
                        'total_categories': len(analysis.data)
                    },
                    confidence="高",
                    priority="高"
                )
                insights.append(insight)
                print(f"[关键洞察] 添加洞察: {insight.title}")
        
        # 2. 从业务指标中提取洞察
        if self.business_metrics:
            # 找出最大值和最小值
            sum_metrics = [m for m in self.business_metrics if m.aggregation_type == 'sum']
            if sum_metrics:
                max_metric = max(sum_metrics, key=lambda x: x.value)
                insight = KeyInsight(
                    title=f"核心业务指标：{max_metric.name}",
                    description=f"{max_metric.name}达到{max_metric.value:.2f}{max_metric.unit}，"
                               f"是数据中最核心的业务指标。",
                    insight_type="pattern",
                    supporting_data={
                        'metric_name': max_metric.name,
                        'metric_value': max_metric.value,
                        'unit': max_metric.unit
                    },
                    confidence="高",
                    priority="中"
                )
                insights.append(insight)
        
        # 3. 如果洞察太少，添加通用洞察
        if len(insights) < 2:
            insight = KeyInsight(
                title="数据整体概览",
                description=f"数据集包含{len(self.df)}条记录，"
                           f"涵盖{len(self.numeric_columns)}个数值指标和{len(self.categorical_columns)}个分类维度。"
                           f"建议进一步深入分析各维度的关联关系。",
                insight_type="pattern",
                supporting_data={
                    'total_records': len(self.df),
                    'numeric_columns': self.numeric_columns,
                    'categorical_columns': self.categorical_columns
                },
                confidence="高",
                priority="中"
            )
            insights.append(insight)
        
        self.key_insights = insights
        print(f"[关键洞察] 总共提取 {len(insights)} 个洞察")
        return insights
    
    def _generate_recommendations(self) -> List[Recommendation]:
        """生成业务建议"""
        recommendations = []
        
        # 基于关键洞察生成建议
        for insight in self.key_insights[:3]:
            if insight.insight_type == "opportunity":
                rec = Recommendation(
                    title=f"把握{insight.title.split('：')[0]}机会",
                    description=f"基于数据分析，{insight.description}"
                               f"建议继续保持该优势，并考虑扩大投入以获取更大市场份额。",
                    recommendation_type="opportunity",
                    expected_impact="提升业务规模，增加收益",
                    implementation_difficulty="中"
                )
                recommendations.append(rec)
            elif insight.insight_type == "risk":
                rec = Recommendation(
                    title=f"关注{insight.title.split('：')[0]}风险",
                    description=f"{insight.description}建议制定应对策略，降低潜在风险。",
                    recommendation_type="improvement",
                    expected_impact="降低风险，稳定业务",
                    implementation_difficulty="中"
                )
                recommendations.append(rec)
        
        # 如果没有建议，添加通用建议
        if not recommendations:
            rec = Recommendation(
                title="持续监控业务指标",
                description="建议建立定期数据分析机制，持续监控关键业务指标的变化趋势，"
                           "及时发现业务机会和潜在风险。",
                recommendation_type="optimization",
                expected_impact="提升决策质量",
                implementation_difficulty="低"
            )
            recommendations.append(rec)
        
        self.recommendations = recommendations
        return recommendations
    
    def _format_metric_table(self, metrics: List[BusinessMetric]) -> pd.DataFrame:
        """格式化业务指标表"""
        data = []
        for m in metrics:
            data.append({
                '指标名称': m.name,
                '指标值': m.format_str.format(m.value),
                '单位': m.unit,
                '指标说明': m.description
            })
        return pd.DataFrame(data)
    
    def _format_dimension_table(self, analysis: DimensionAnalysis) -> pd.DataFrame:
        """格式化维度分析表"""
        return analysis.data.copy()
    
    def _format_table_markdown(self, df: pd.DataFrame, title: str = "") -> str:
        """格式化Markdown表格"""
        lines = []
        if title:
            lines.append(f"**{title}**")
            lines.append("")
        
        # 表头
        headers = df.columns.tolist()
        lines.append("| " + " | ".join(headers) + " |")
        lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
        
        # 数据行
        for _, row in df.iterrows():
            row_values = [str(v) for v in row.values]
            lines.append("| " + " | ".join(row_values) + " |")
        
        lines.append("")
        return "\n".join(lines)
    
    def generate_markdown_report(self, title: str = "业务分析报告") -> str:
        """生成Markdown报告 - 改进版本"""
        
        # 执行分析
        print("[报告生成] 开始执行分析...")
        self._analyze_business_metrics()
        self._analyze_by_dimensions()
        self._extract_key_insights()
        self._generate_recommendations()
        
        print(f"[报告生成] 指标: {len(self.business_metrics)}, 维度: {len(self.dimension_analyses)}, 洞察: {len(self.key_insights)}")
        
        report_lines = []
        
        # 报告标题
        report_lines.append(f"# {title}")
        report_lines.append("")
        report_lines.append(f"**生成时间**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        report_lines.append("")
        report_lines.append("---")
        report_lines.append("")
        
        # 章节1：核心业务指标
        report_lines.append("## 一、核心业务指标")
        report_lines.append("")
        report_lines.append("本章节展示了数据中的核心业务指标汇总统计，反映整体业务规模。")
        report_lines.append("")
        
        if self.business_metrics:
            metric_table = self._format_metric_table(self.business_metrics)
            report_lines.append(self._format_table_markdown(metric_table, "核心业务指标统计表"))
            
            # 添加指标说明
            report_lines.append("**指标说明**：")
            report_lines.append("- 汇总：该指标所有数据的总和，反映整体规模")
            report_lines.append("- 平均：该指标的平均值，反映一般水平")
            report_lines.append("- 最大/最小：该指标的极值，反映波动范围")
            report_lines.append("")
        else:
            report_lines.append("*未识别到数值型业务指标*")
            report_lines.append("")
        
        report_lines.append("---")
        report_lines.append("")
        
        # 章节2：维度分析
        report_lines.append("## 二、维度分析")
        report_lines.append("")
        report_lines.append("本章节按不同维度对业务指标进行分析，识别各维度的表现差异和业务重点。")
        report_lines.append("")
        
        if self.dimension_analyses:
            for i, analysis in enumerate(self.dimension_analyses[:5], 1):
                report_lines.append(f"### 2.{i} {analysis.dimension_name}维度 - {analysis.metric_name}")
                report_lines.append("")
                
                # 添加分析总结
                report_lines.append(f"**分析总结**：{analysis.summary}")
                report_lines.append("")
                
                # 添加表格
                dim_table = self._format_dimension_table(analysis)
                report_lines.append(self._format_table_markdown(dim_table, f"按{analysis.dimension_name}分析"))
                
                # 添加详细说明
                if analysis.top_values:
                    top1 = analysis.top_values[0]
                    report_lines.append(f"**详细说明**：")
                    report_lines.append(f"- 表现最优：{top1.get(analysis.dimension_name, '未知')}，"
                                       f"汇总值{top1.get('汇总', 0):.2f}")
                    if len(analysis.top_values) > 1:
                        top2 = analysis.top_values[1]
                        report_lines.append(f"- 排名第二：{top2.get(analysis.dimension_name, '未知')}，"
                                           f"汇总值{top2.get('汇总', 0):.2f}")
                    report_lines.append(f"- 数据覆盖：共{len(analysis.data)}个类别，"
                                       f"平均每个类别{analysis.data['平均'].mean():.2f}")
                    report_lines.append("")
                
                report_lines.append("")
        else:
            report_lines.append("*未识别到有效的维度分析*")
            report_lines.append("")
        
        report_lines.append("---")
        report_lines.append("")
        
        # 章节3：关键洞察
        report_lines.append("## 三、关键洞察")
        report_lines.append("")
        report_lines.append("本章节基于数据分析提取关键业务洞察，为决策提供数据支撑。")
        report_lines.append("")
        
        if self.key_insights:
            for i, insight in enumerate(self.key_insights, 1):
                type_names = {
                    'anomaly': '⚠️ 异常发现',
                    'opportunity': '💡 业务机会',
                    'pattern': '📊 数据模式',
                    'risk': '⚡ 潜在风险'
                }
                type_name = type_names.get(insight.insight_type, insight.insight_type)
                
                report_lines.append(f"### 3.{i} {insight.title}")
                report_lines.append("")
                report_lines.append(f"**洞察类型**：{type_name}")
                report_lines.append("")
                report_lines.append(f"**洞察描述**：{insight.description}")
                report_lines.append("")
                
                # 支撑数据表格
                if insight.supporting_data:
                    report_lines.append("**支撑数据**：")
                    support_df = pd.DataFrame([insight.supporting_data])
                    report_lines.append(self._format_table_markdown(support_df))
                
                report_lines.append(f"*置信度：{insight.confidence} | 优先级：{insight.priority}*")
                report_lines.append("")
        else:
            report_lines.append("*暂无关键洞察*")
            report_lines.append("")
        
        report_lines.append("---")
        report_lines.append("")
        
        # 章节4：业务建议
        report_lines.append("## 四、业务建议")
        report_lines.append("")
        report_lines.append("本章节基于关键洞察提出可执行的业务建议。")
        report_lines.append("")
        
        if self.recommendations:
            for i, rec in enumerate(self.recommendations, 1):
                type_names = {
                    'improvement': '🔧 改进建议',
                    'opportunity': '🚀 机会建议',
                    'optimization': '⚙️ 优化建议'
                }
                type_name = type_names.get(rec.recommendation_type, rec.recommendation_type)
                
                report_lines.append(f"### 4.{i} {rec.title}")
                report_lines.append("")
                report_lines.append(f"**建议类型**：{type_name}")
                report_lines.append("")
                report_lines.append(f"**建议内容**：{rec.description}")
                report_lines.append("")
                
                if rec.expected_impact:
                    report_lines.append(f"**预期效果**：{rec.expected_impact}")
                    report_lines.append("")
                
                report_lines.append(f"*实施难度：{rec.implementation_difficulty}*")
                report_lines.append("")
        else:
            report_lines.append("*暂无业务建议*")
            report_lines.append("")
        
        report_lines.append("---")
        report_lines.append("")
        
        # 报告结尾
        report_lines.append(f"*报告生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}*")
        
        return "\n".join(report_lines)
    
    def generate_word_report(self, title: str = "业务分析报告") -> io.BytesIO:
        """生成Word报告"""
        # 执行分析
        self._analyze_business_metrics()
        self._analyze_by_dimensions()
        self._extract_key_insights()
        self._generate_recommendations()
        
        doc = Document()
        
        # 设置默认字体
        style = doc.styles['Normal']
        style.font.name = 'Microsoft YaHei'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        
        # 标题
        title_heading = doc.add_heading(title, 0)
        for run in title_heading.runs:
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            run.font.size = Pt(18)
            run.font.bold = True
        title_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 元信息
        p = doc.add_paragraph()
        run = p.add_run(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        
        doc.add_paragraph()
        
        # 章节1：核心业务指标
        heading1 = doc.add_heading('一、核心业务指标', level=1)
        for run in heading1.runs:
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            run.font.size = Pt(14)
            run.font.bold = True
        
        p = doc.add_paragraph()
        run = p.add_run("本章节展示了数据中的核心业务指标汇总统计，反映整体业务规模。")
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        run.font.size = Pt(10)
        doc.add_paragraph()
        
        if self.business_metrics:
            # 添加表格
            metric_table = self._format_metric_table(self.business_metrics)
            table = doc.add_table(rows=1, cols=len(metric_table.columns))
            table.style = 'Light Grid Accent 1'
            
            # 表头
            hdr_cells = table.rows[0].cells
            for i, col in enumerate(metric_table.columns):
                hdr_cells[i].text = col
                for paragraph in hdr_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.name = 'Microsoft YaHei'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            
            # 数据行
            for _, row in metric_table.iterrows():
                row_cells = table.add_row().cells
                for i, value in enumerate(row):
                    row_cells[i].text = str(value)
                    for paragraph in row_cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Microsoft YaHei'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        
        doc.add_page_break()
        
        # 保存文档
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output


def generate_business_semantic_report(df: pd.DataFrame, title: str = "业务分析报告") -> Tuple[str, io.BytesIO]:
    """生成业务语义分析报告的便捷函数"""
    generator = BusinessSemanticReportGeneratorV2(df)
    
    md_report = generator.generate_markdown_report(title=title)
    word_report = generator.generate_word_report(title=title)
    
    return md_report, word_report

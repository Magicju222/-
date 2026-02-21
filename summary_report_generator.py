"""
Summary Report Generator
生成数据分析报告的Markdown和Word文档
"""

import json
import re
from typing import Any, Dict, List, Optional, Union
from datetime import datetime

import pandas as pd
import numpy as np
from scipy import stats


# ==================== 数据提取函数 ====================

def _extract_all_step_data(agent_result) -> List[Dict[str, Any]]:
    """
    从所有Agent步骤中提取完整数据
    
    Args:
        agent_result: Agent执行结果对象
        
    Returns:
        包含所有步骤数据的列表
    """
    steps_data = []
    
    if not hasattr(agent_result, 'steps') or not agent_result.steps:
        return steps_data
    
    for step in agent_result.steps:
        step_info = {
            'step_number': getattr(step, 'step_number', None),
            'action': getattr(step, 'action', None),
            'thought': getattr(step, 'thought', None),
            'observation': getattr(step, 'observation', None),
            'tool_result': None,
            'parsed_data': None
        }
        
        # 提取tool_result
        tool_result = getattr(step, 'tool_result', None)
        if tool_result:
            step_info['tool_result'] = tool_result
            
            # 如果tool_result成功，尝试解析结果为表格格式
            if isinstance(tool_result, dict) and tool_result.get('success'):
                result_data = tool_result.get('result') or tool_result.get('data')
                if result_data is not None:
                    step_info['parsed_data'] = _parse_data_to_table(result_data)
        
        steps_data.append(step_info)
    
    return steps_data


def _parse_data_to_table(data: Any) -> Dict[str, Any]:
    """
    将数据解析为表格格式
    
    Args:
        data: 输入数据（字典、列表等）
        
    Returns:
        {'headers': [], 'rows': []} 格式的表格数据
    """
    table_data = {'headers': [], 'rows': []}
    
    if data is None:
        return table_data
    
    # 处理pandas DataFrame
    if isinstance(data, pd.DataFrame):
        table_data['headers'] = data.columns.tolist()
        table_data['rows'] = data.values.tolist()
        return table_data
    
    # 处理字典类型数据
    if isinstance(data, dict):
        # describe() 结果
        if all(isinstance(v, (dict, pd.Series)) for v in data.values()):
            # 转置describe结果
            headers = ['字段'] + list(list(data.values())[0].keys() if isinstance(list(data.values())[0], dict) else list(data.values())[0].index)
            rows = []
            for key, value in data.items():
                if isinstance(value, dict):
                    row = [key] + [str(v) for v in value.values()]
                elif isinstance(value, pd.Series):
                    row = [key] + [str(v) for v in value.values]
                else:
                    row = [key, str(value)]
                rows.append(row)
            table_data['headers'] = headers
            table_data['rows'] = rows
        # head() 结果或普通字典
        else:
            table_data['headers'] = list(data.keys())
            table_data['rows'] = [list(data.values())]
        return table_data
    
    # 处理列表类型数据
    if isinstance(data, list):
        if len(data) == 0:
            return table_data
        
        # 列表中的字典（记录列表）
        if all(isinstance(item, dict) for item in data):
            # 收集所有可能的字段
            all_keys = set()
            for item in data:
                all_keys.update(item.keys())
            table_data['headers'] = sorted(list(all_keys))
            for item in data:
                row = [str(item.get(k, '')) for k in table_data['headers']]
                table_data['rows'].append(row)
        # 简单列表
        else:
            table_data['headers'] = ['值']
            table_data['rows'] = [[str(item)] for item in data]
        return table_data
    
    # 处理pandas Series
    if isinstance(data, pd.Series):
        table_data['headers'] = [data.name or '值']
        table_data['rows'] = [[str(v)] for v in data.values]
        return table_data
    
    # 其他类型
    table_data['headers'] = ['值']
    table_data['rows'] = [[str(data)]]
    return table_data


def _format_table_markdown(table_data: Dict[str, List]) -> str:
    """
    将表格格式化为Markdown
    
    Args:
        table_data: {'headers': [], 'rows': []} 格式的表格数据
        
    Returns:
        Markdown格式的表格字符串
    """
    if not table_data or not table_data.get('headers'):
        return ''
    
    headers = table_data['headers']
    rows = table_data.get('rows', [])
    
    # 生成表头
    header_line = '| ' + ' | '.join(str(h) for h in headers) + ' |'
    
    # 生成分隔符
    separator_line = '|' + '|'.join([' --- ' for _ in headers]) + '|'
    
    # 生成数据行
    row_lines = []
    for row in rows:
        # 确保行数据与表头数量一致
        row_values = [str(row[i]) if i < len(row) else '' for i in range(len(headers))]
        row_lines.append('| ' + ' | '.join(row_values) + ' |')
    
    # 组合所有行
    markdown_lines = [header_line, separator_line] + row_lines
    return '\n'.join(markdown_lines)


# ==================== 分析函数 ====================

def _analyze_business_metrics(dfs: Dict[str, pd.DataFrame]) -> List[Dict[str, Any]]:
    """
    分析业务指标
    
    Args:
        dfs: 包含DataFrame的字典
        
    Returns:
        指标列表，每个指标包含名称、汇总值、平均值、最大值、最小值、标准差
    """
    metrics = []
    
    for sheet_name, df in dfs.items():
        if df is None or df.empty:
            continue
        
        # 识别所有数值型字段
        numeric_columns = df.select_dtypes(include=[np.number]).columns.tolist()
        
        for col in numeric_columns:
            # 过滤空值和零值
            values = df[col].dropna()
            values = values[values != 0]
            
            if len(values) == 0:
                continue
            
            metric = {
                'sheet_name': sheet_name,
                'field_name': col,
                'count': int(len(values)),
                'sum': float(values.sum()),
                'mean': float(values.mean()),
                'max': float(values.max()),
                'min': float(values.min()),
                'std': float(values.std()) if len(values) > 1 else 0.0
            }
            metrics.append(metric)
    
    return metrics


def _analyze_categories(dfs: Dict[str, pd.DataFrame]) -> List[Dict[str, Any]]:
    """
    分析分类数据
    
    Args:
        dfs: 包含DataFrame的字典
        
    Returns:
        分类列表，每个分类包含字段名、类别数量、分布情况
    """
    categories = []
    
    for sheet_name, df in dfs.items():
        if df is None or df.empty:
            continue
        
        # 识别所有分类字段（对象类型或类别类型）
        categorical_columns = df.select_dtypes(include=['object', 'category']).columns.tolist()
        
        for col in categorical_columns:
            # 过滤空值、空字符串、'0'
            values = df[col].dropna()
            values = values[values != '']
            values = values[values != '0']
            values = values[values.astype(str).str.strip() != '']
            
            if len(values) == 0:
                continue
            
            # 统计类别分布
            value_counts = values.value_counts()
            
            category = {
                'sheet_name': sheet_name,
                'field_name': col,
                'unique_count': int(len(value_counts)),
                'total_count': int(len(values)),
                'distribution': value_counts.head(10).to_dict(),
                'top_category': value_counts.index[0] if len(value_counts) > 0 else None,
                'top_count': int(value_counts.iloc[0]) if len(value_counts) > 0 else 0
            }
            categories.append(category)
    
    return categories


def _detect_trends(dfs: Dict[str, pd.DataFrame]) -> List[Dict[str, Any]]:
    """
    检测趋势
    
    Args:
        dfs: 包含DataFrame的字典
        
    Returns:
        趋势列表，每个趋势包含字段名、斜率、趋势方向
    """
    trends = []
    
    for sheet_name, df in dfs.items():
        if df is None or df.empty:
            continue
        
        # 识别数值型字段
        numeric_columns = df.select_dtypes(include=[np.number]).columns.tolist()
        
        for col in numeric_columns:
            values = df[col].dropna()
            
            if len(values) < 3:  # 需要至少3个点才能检测趋势
                continue
            
            # 使用线性回归计算斜率
            x = np.arange(len(values))
            slope, intercept, r_value, p_value, std_err = stats.linregress(x, values)
            
            # 识别趋势方向
            if slope > 0.01:
                trend_direction = '上升'
            elif slope < -0.01:
                trend_direction = '下降'
            else:
                trend_direction = '平稳'
            
            trend = {
                'sheet_name': sheet_name,
                'field_name': col,
                'slope': float(slope),
                'r_squared': float(r_value ** 2),
                'p_value': float(p_value),
                'trend_direction': trend_direction,
                'significance': '显著' if p_value < 0.05 else '不显著'
            }
            trends.append(trend)
    
    return trends


def _detect_anomalies(dfs: Dict[str, pd.DataFrame]) -> List[Dict[str, Any]]:
    """
    检测异常
    
    Args:
        dfs: 包含DataFrame的字典
        
    Returns:
        异常列表，每个异常包含字段名、异常数量、异常比例
    """
    anomalies = []
    
    for sheet_name, df in dfs.items():
        if df is None or df.empty:
            continue
        
        # 识别数值型字段
        numeric_columns = df.select_dtypes(include=[np.number]).columns.tolist()
        
        for col in numeric_columns:
            values = df[col].dropna()
            
            if len(values) < 3:
                continue
            
            # 使用3σ原则检测异常值
            mean = values.mean()
            std = values.std()
            
            if std == 0:
                continue
            
            # 计算z-score
            z_scores = np.abs((values - mean) / std)
            
            # 识别异常值（|z-score| > 3）
            anomaly_mask = z_scores > 3
            anomaly_count = int(anomaly_mask.sum())
            anomaly_ratio = float(anomaly_count / len(values))
            
            if anomaly_count > 0:
                anomaly_values = values[anomaly_mask].tolist()
            else:
                anomaly_values = []
            
            anomaly = {
                'sheet_name': sheet_name,
                'field_name': col,
                'total_count': int(len(values)),
                'anomaly_count': anomaly_count,
                'anomaly_ratio': anomaly_ratio,
                'mean': float(mean),
                'std': float(std),
                'anomaly_values': anomaly_values[:10]  # 最多显示10个异常值
            }
            anomalies.append(anomaly)
    
    return anomalies


# ==================== 报告生成函数 ====================

def generate_markdown_report(agent_result, dfs: Dict[str, pd.DataFrame], sheet_name: str) -> str:
    """
    生成Markdown报告
    
    Args:
        agent_result: Agent执行结果对象
        dfs: 包含DataFrame的字典
        sheet_name: 工作表名称
        
    Returns:
        Markdown格式的报告字符串
    """
    report_lines = []
    
    # 报告标题
    report_lines.append(f'数据分析报告 - {sheet_name}')
    report_lines.append('=' * 50)
    report_lines.append('')
    
    # 一、业务指标汇总
    report_lines.append('一、业务指标汇总')
    report_lines.append('-' * 30)
    report_lines.append('')
    
    metrics = _analyze_business_metrics(dfs)
    if metrics:
        # 创建指标汇总表
        metric_table = {
            'headers': ['工作表', '字段', '记录数', '汇总值', '平均值', '最大值', '最小值', '标准差'],
            'rows': []
        }
        for m in metrics:
            metric_table['rows'].append([
                m['sheet_name'],
                m['field_name'],
                m['count'],
                f"{m['sum']:.2f}",
                f"{m['mean']:.2f}",
                f"{m['max']:.2f}",
                f"{m['min']:.2f}",
                f"{m['std']:.2f}"
            ])
        report_lines.append(_format_table_markdown(metric_table))
        report_lines.append('')
    else:
        report_lines.append('未找到数值型业务指标。')
        report_lines.append('')
    
    # 二、详细分析过程
    report_lines.append('二、详细分析过程')
    report_lines.append('-' * 30)
    report_lines.append('')
    
    # 提取所有步骤数据
    steps_data = _extract_all_step_data(agent_result)
    
    if steps_data:
        for step in steps_data:
            step_num = step.get('step_number', 'N/A')
            report_lines.append(f'步骤 {step_num}')
            report_lines.append('')
            
            # 显示动作
            action = step.get('action')
            if action:
                report_lines.append(f'执行动作: {action}')
                report_lines.append('')
            
            # 显示思考过程
            thought = step.get('thought')
            if thought:
                report_lines.append('分析思路:')
                report_lines.append(thought)
                report_lines.append('')
            
            # 显示解析后的数据表格
            parsed_data = step.get('parsed_data')
            if parsed_data and parsed_data.get('headers'):
                report_lines.append('数据结果:')
                report_lines.append(_format_table_markdown(parsed_data))
                report_lines.append('')
            
            # 显示观察结果
            observation = step.get('observation')
            if observation:
                report_lines.append('观察结果:')
                report_lines.append(observation)
                report_lines.append('')
            
            report_lines.append('---')
            report_lines.append('')
    else:
        report_lines.append('暂无详细分析步骤数据。')
        report_lines.append('')
    
    # 三、业务洞察总结
    report_lines.append('三、业务洞察总结')
    report_lines.append('-' * 30)
    report_lines.append('')
    
    # 分类分析
    report_lines.append('1. 分类数据分析')
    report_lines.append('')
    
    categories = _analyze_categories(dfs)
    if categories:
        cat_table = {
            'headers': ['工作表', '字段', '类别数', '总记录数', '主要类别', '主要类别数量'],
            'rows': []
        }
        for c in categories:
            cat_table['rows'].append([
                c['sheet_name'],
                c['field_name'],
                c['unique_count'],
                c['total_count'],
                str(c['top_category']),
                c['top_count']
            ])
        report_lines.append(_format_table_markdown(cat_table))
        report_lines.append('')
    else:
        report_lines.append('未找到分类数据。')
        report_lines.append('')
    
    # 趋势分析
    report_lines.append('2. 趋势分析')
    report_lines.append('')
    
    trends = _detect_trends(dfs)
    if trends:
        trend_table = {
            'headers': ['工作表', '字段', '趋势方向', '斜率', 'R方', '显著性'],
            'rows': []
        }
        for t in trends:
            trend_table['rows'].append([
                t['sheet_name'],
                t['field_name'],
                t['trend_direction'],
                f"{t['slope']:.4f}",
                f"{t['r_squared']:.4f}",
                t['significance']
            ])
        report_lines.append(_format_table_markdown(trend_table))
        report_lines.append('')
    else:
        report_lines.append('未检测到明显趋势。')
        report_lines.append('')
    
    # 异常检测
    report_lines.append('3. 异常检测')
    report_lines.append('')
    
    anomalies = _detect_anomalies(dfs)
    if anomalies:
        # 只显示有异常的字段
        anomalies_with_data = [a for a in anomalies if a['anomaly_count'] > 0]
        if anomalies_with_data:
            anomaly_table = {
                'headers': ['工作表', '字段', '异常数量', '异常比例', '异常值示例'],
                'rows': []
            }
            for a in anomalies_with_data:
                anomaly_values_str = ', '.join([f"{v:.2f}" for v in a['anomaly_values'][:5]])
                if len(a['anomaly_values']) > 5:
                    anomaly_values_str += '...'
                anomaly_table['rows'].append([
                    a['sheet_name'],
                    a['field_name'],
                    a['anomaly_count'],
                    f"{a['anomaly_ratio']*100:.2f}%",
                    anomaly_values_str
                ])
            report_lines.append(_format_table_markdown(anomaly_table))
            report_lines.append('')
        else:
            report_lines.append('未检测到异常值（使用3σ原则）。')
            report_lines.append('')
    else:
        report_lines.append('未进行异常检测。')
        report_lines.append('')
    
    # 四、结论与建议
    report_lines.append('四、结论与建议')
    report_lines.append('-' * 30)
    report_lines.append('')
    
    # 自动生成结论
    conclusions = []
    
    if metrics:
        total_fields = len(set(m['field_name'] for m in metrics))
        conclusions.append(f'数据包含 {total_fields} 个数值型字段，可进行量化分析。')
    
    if categories:
        total_cat_fields = len(set(c['field_name'] for c in categories))
        conclusions.append(f'数据包含 {total_cat_fields} 个分类字段，可用于分组分析。')
    
    if trends:
        up_trends = [t for t in trends if t['trend_direction'] == '上升']
        down_trends = [t for t in trends if t['trend_direction'] == '下降']
        if up_trends:
            conclusions.append(f'检测到 {len(up_trends)} 个指标呈上升趋势。')
        if down_trends:
            conclusions.append(f'检测到 {len(down_trends)} 个指标呈下降趋势。')
    
    if anomalies:
        total_anomalies = sum(a['anomaly_count'] for a in anomalies)
        if total_anomalies > 0:
            conclusions.append(f'检测到 {total_anomalies} 个异常值，建议进一步核查。')
        else:
            conclusions.append('数据质量良好，未检测到明显异常值。')
    
    if conclusions:
        for i, conclusion in enumerate(conclusions, 1):
            report_lines.append(f'{i}. {conclusion}')
        report_lines.append('')
    
    report_lines.append('建议:')
    report_lines.append('')
    report_lines.append('1. 定期监控关键业务指标的变化趋势。')
    report_lines.append('2. 对异常数据进行深入调查，找出根本原因。')
    report_lines.append('3. 基于分类数据特征，制定针对性的业务策略。')
    report_lines.append('4. 持续优化数据质量，确保分析结果准确可靠。')
    report_lines.append('')
    
    # 报告生成时间
    report_lines.append('')
    report_lines.append('-' * 30)
    report_lines.append(f'报告生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    
    return '\n'.join(report_lines)


def generate_word_report(agent_result, dfs: Dict[str, pd.DataFrame], sheet_name: str) -> str:
    """
    生成Word报告
    
    Args:
        agent_result: Agent执行结果对象
        dfs: 包含DataFrame的字典
        sheet_name: 工作表名称
        
    Returns:
        生成的Word文件路径
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
    except ImportError:
        raise ImportError("请安装 python-docx 库: pip install python-docx")
    
    # 创建文档
    doc = Document()
    
    # 设置中文字体
    def set_chinese_font(run, font_name='Microsoft YaHei', font_size=11, bold=False):
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run.font.size = Pt(font_size)
        run.font.bold = bold
    
    # 添加标题
    title = doc.add_heading(f'数据分析报告 - {sheet_name}', level=0)
    for run in title.runs:
        set_chinese_font(run, font_size=18, bold=True)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # 一、业务指标汇总
    heading1 = doc.add_heading('一、业务指标汇总', level=1)
    for run in heading1.runs:
        set_chinese_font(run, font_size=14, bold=True)
    
    metrics = _analyze_business_metrics(dfs)
    if metrics:
        # 创建指标汇总表
        table = doc.add_table(rows=1, cols=8)
        table.style = 'Light Grid Accent 1'
        
        # 表头
        headers = ['工作表', '字段', '记录数', '汇总值', '平均值', '最大值', '最小值', '标准差']
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    set_chinese_font(run, bold=True)
        
        # 数据行
        for m in metrics:
            row_cells = table.add_row().cells
            row_cells[0].text = m['sheet_name']
            row_cells[1].text = m['field_name']
            row_cells[2].text = str(m['count'])
            row_cells[3].text = f"{m['sum']:.2f}"
            row_cells[4].text = f"{m['mean']:.2f}"
            row_cells[5].text = f"{m['max']:.2f}"
            row_cells[6].text = f"{m['min']:.2f}"
            row_cells[7].text = f"{m['std']:.2f}"
            
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_chinese_font(run)
        
        doc.add_paragraph()
    else:
        p = doc.add_paragraph('未找到数值型业务指标。')
        for run in p.runs:
            set_chinese_font(run)
    
    # 二、详细分析过程
    heading2 = doc.add_heading('二、详细分析过程', level=1)
    for run in heading2.runs:
        set_chinese_font(run, font_size=14, bold=True)
    
    steps_data = _extract_all_step_data(agent_result)
    
    if steps_data:
        for step in steps_data:
            step_num = step.get('step_number', 'N/A')
            
            # 步骤标题
            step_heading = doc.add_heading(f'步骤 {step_num}', level=2)
            for run in step_heading.runs:
                set_chinese_font(run, font_size=12, bold=True)
            
            # 执行动作
            action = step.get('action')
            if action:
                p = doc.add_paragraph()
                run = p.add_run(f'执行动作: ')
                set_chinese_font(run, bold=True)
                run = p.add_run(action)
                set_chinese_font(run)
            
            # 分析思路
            thought = step.get('thought')
            if thought:
                p = doc.add_paragraph()
                run = p.add_run('分析思路:')
                set_chinese_font(run, bold=True)
                p = doc.add_paragraph(thought)
                for run in p.runs:
                    set_chinese_font(run)
            
            # 数据结果表格
            parsed_data = step.get('parsed_data')
            if parsed_data and parsed_data.get('headers'):
                p = doc.add_paragraph()
                run = p.add_run('数据结果:')
                set_chinese_font(run, bold=True)
                
                headers = parsed_data['headers']
                rows = parsed_data['rows']
                
                # 限制行数，避免表格过大
                display_rows = rows[:20]
                if len(rows) > 20:
                    display_rows.append(['...'] * len(headers))
                
                table = doc.add_table(rows=1, cols=len(headers))
                table.style = 'Light List Accent 1'
                
                # 表头
                hdr_cells = table.rows[0].cells
                for i, header in enumerate(headers):
                    hdr_cells[i].text = str(header)
                    for paragraph in hdr_cells[i].paragraphs:
                        for run in paragraph.runs:
                            set_chinese_font(run, bold=True)
                
                # 数据行
                for row in display_rows:
                    row_cells = table.add_row().cells
                    for i, value in enumerate(row):
                        if i < len(row_cells):
                            row_cells[i].text = str(value)
                            for paragraph in row_cells[i].paragraphs:
                                for run in paragraph.runs:
                                    set_chinese_font(run)
            
            # 观察结果
            observation = step.get('observation')
            if observation:
                p = doc.add_paragraph()
                run = p.add_run('观察结果:')
                set_chinese_font(run, bold=True)
                p = doc.add_paragraph(observation)
                for run in p.runs:
                    set_chinese_font(run)
            
            doc.add_paragraph()
    else:
        p = doc.add_paragraph('暂无详细分析步骤数据。')
        for run in p.runs:
            set_chinese_font(run)
    
    # 三、业务洞察总结
    heading3 = doc.add_heading('三、业务洞察总结', level=1)
    for run in heading3.runs:
        set_chinese_font(run, font_size=14, bold=True)
    
    # 分类分析
    cat_heading = doc.add_heading('1. 分类数据分析', level=2)
    for run in cat_heading.runs:
        set_chinese_font(run, font_size=12, bold=True)
    
    categories = _analyze_categories(dfs)
    if categories:
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Light Grid Accent 1'
        
        headers = ['工作表', '字段', '类别数', '总记录数', '主要类别', '主要类别数量']
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    set_chinese_font(run, bold=True)
        
        for c in categories:
            row_cells = table.add_row().cells
            row_cells[0].text = c['sheet_name']
            row_cells[1].text = c['field_name']
            row_cells[2].text = str(c['unique_count'])
            row_cells[3].text = str(c['total_count'])
            row_cells[4].text = str(c['top_category'])
            row_cells[5].text = str(c['top_count'])
            
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_chinese_font(run)
        
        doc.add_paragraph()
    else:
        p = doc.add_paragraph('未找到分类数据。')
        for run in p.runs:
            set_chinese_font(run)
    
    # 趋势分析
    trend_heading = doc.add_heading('2. 趋势分析', level=2)
    for run in trend_heading.runs:
        set_chinese_font(run, font_size=12, bold=True)
    
    trends = _detect_trends(dfs)
    if trends:
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Light Grid Accent 1'
        
        headers = ['工作表', '字段', '趋势方向', '斜率', 'R方', '显著性']
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    set_chinese_font(run, bold=True)
        
        for t in trends:
            row_cells = table.add_row().cells
            row_cells[0].text = t['sheet_name']
            row_cells[1].text = t['field_name']
            row_cells[2].text = t['trend_direction']
            row_cells[3].text = f"{t['slope']:.4f}"
            row_cells[4].text = f"{t['r_squared']:.4f}"
            row_cells[5].text = t['significance']
            
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_chinese_font(run)
        
        doc.add_paragraph()
    else:
        p = doc.add_paragraph('未检测到明显趋势。')
        for run in p.runs:
            set_chinese_font(run)
    
    # 异常检测
    anomaly_heading = doc.add_heading('3. 异常检测', level=2)
    for run in anomaly_heading.runs:
        set_chinese_font(run, font_size=12, bold=True)
    
    anomalies = _detect_anomalies(dfs)
    if anomalies:
        anomalies_with_data = [a for a in anomalies if a['anomaly_count'] > 0]
        if anomalies_with_data:
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Light Grid Accent 1'
            
            headers = ['工作表', '字段', '异常数量', '异常比例', '异常值示例']
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                for paragraph in hdr_cells[i].paragraphs:
                    for run in paragraph.runs:
                        set_chinese_font(run, bold=True)
            
            for a in anomalies_with_data:
                row_cells = table.add_row().cells
                anomaly_values_str = ', '.join([f"{v:.2f}" for v in a['anomaly_values'][:5]])
                if len(a['anomaly_values']) > 5:
                    anomaly_values_str += '...'
                
                row_cells[0].text = a['sheet_name']
                row_cells[1].text = a['field_name']
                row_cells[2].text = str(a['anomaly_count'])
                row_cells[3].text = f"{a['anomaly_ratio']*100:.2f}%"
                row_cells[4].text = anomaly_values_str
                
                for cell in row_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            set_chinese_font(run)
            
            doc.add_paragraph()
        else:
            p = doc.add_paragraph('未检测到异常值（使用3σ原则）。')
            for run in p.runs:
                set_chinese_font(run)
    else:
        p = doc.add_paragraph('未进行异常检测。')
        for run in p.runs:
            set_chinese_font(run)
    
    # 四、结论与建议
    heading4 = doc.add_heading('四、结论与建议', level=1)
    for run in heading4.runs:
        set_chinese_font(run, font_size=14, bold=True)
    
    # 自动生成结论
    conclusions = []
    
    if metrics:
        total_fields = len(set(m['field_name'] for m in metrics))
        conclusions.append(f'数据包含 {total_fields} 个数值型字段，可进行量化分析。')
    
    if categories:
        total_cat_fields = len(set(c['field_name'] for c in categories))
        conclusions.append(f'数据包含 {total_cat_fields} 个分类字段，可用于分组分析。')
    
    if trends:
        up_trends = [t for t in trends if t['trend_direction'] == '上升']
        down_trends = [t for t in trends if t['trend_direction'] == '下降']
        if up_trends:
            conclusions.append(f'检测到 {len(up_trends)} 个指标呈上升趋势。')
        if down_trends:
            conclusions.append(f'检测到 {len(down_trends)} 个指标呈下降趋势。')
    
    if anomalies:
        total_anomalies = sum(a['anomaly_count'] for a in anomalies)
        if total_anomalies > 0:
            conclusions.append(f'检测到 {total_anomalies} 个异常值，建议进一步核查。')
        else:
            conclusions.append('数据质量良好，未检测到明显异常值。')
    
    if conclusions:
        for i, conclusion in enumerate(conclusions, 1):
            p = doc.add_paragraph(f'{i}. {conclusion}')
            for run in p.runs:
                set_chinese_font(run)
        doc.add_paragraph()
    
    # 建议
    p = doc.add_paragraph()
    run = p.add_run('建议:')
    set_chinese_font(run, bold=True)
    
    suggestions = [
        '1. 定期监控关键业务指标的变化趋势。',
        '2. 对异常数据进行深入调查，找出根本原因。',
        '3. 基于分类数据特征，制定针对性的业务策略。',
        '4. 持续优化数据质量，确保分析结果准确可靠。'
    ]
    
    for suggestion in suggestions:
        p = doc.add_paragraph(suggestion)
        for run in p.runs:
            set_chinese_font(run)
    
    doc.add_paragraph()
    
    # 报告生成时间
    p = doc.add_paragraph(f'报告生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    for run in p.runs:
        set_chinese_font(run)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # 保存文档
    output_path = f'e:\\徐衡文档\\AI\\Trae EXCEL\\数据分析报告_{sheet_name}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
    doc.save(output_path)
    
    return output_path

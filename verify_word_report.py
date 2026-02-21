"""
验证Word报告的表格格式
"""
from docx import Document

def verify_word_report():
    docx_path = 'e:\\徐衡文档\\AI\\Trae EXCEL\\test_data\\mock_report_区域业绩.docx'
    
    print("="*80)
    print("Word报告验证")
    print("="*80)
    
    doc = Document(docx_path)
    
    # 统计信息
    print(f"\n1. 文档统计:")
    print(f"   段落数: {len(doc.paragraphs)}")
    print(f"   表格数: {len(doc.tables)}")
    
    # 检查表格
    print(f"\n2. 表格详情:")
    for i, table in enumerate(doc.tables, 1):
        print(f"\n   表格 {i}:")
        print(f"   - 行数: {len(table.rows)}")
        print(f"   - 列数: {len(table.columns)}")
        print(f"   - 表头: {[cell.text for cell in table.rows[0].cells]}")
        print(f"   - 首行数据: {[cell.text for cell in table.rows[1].cells] if len(table.rows) > 1 else '无'}")
    
    # 检查段落标题
    print(f"\n3. 文档结构:")
    headings = []
    for para in doc.paragraphs:
        if para.style.name.startswith('Heading'):
            headings.append(para.text)
    
    for h in headings:
        print(f"   - {h}")
    
    print("\n" + "="*80)
    print("验证完成!")
    print("="*80)

if __name__ == '__main__':
    verify_word_report()

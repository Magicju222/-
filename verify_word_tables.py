"""
验证Word报告中的表格
"""
from docx import Document

def verify_word_tables():
    docx_path = 'e:\\徐衡文档\\AI\\Trae EXCEL\\test_data\\mock_report_区域业绩.docx'
    
    print("="*80)
    print("Word报告表格验证")
    print("="*80)
    
    try:
        doc = Document(docx_path)
        
        # 统计信息
        print(f"\n1. 文档统计:")
        print(f"   段落数: {len(doc.paragraphs)}")
        print(f"   表格数: {len(doc.tables)}")
        
        # 检查表格
        print(f"\n2. 表格详情:")
        for i, table in enumerate(doc.tables, 1):
            print(f"\n   表格 {i}:")
            print(f"   - 行数: {len(table.rows)}")
            print(f"   - 列数: {len(table.columns)}")
            
            # 显示表头
            if len(table.rows) > 0:
                headers = [cell.text for cell in table.rows[0].cells]
                print(f"   - 表头: {headers}")
            
            # 显示前3行数据
            if len(table.rows) > 1:
                print(f"   - 数据预览:")
                for j, row in enumerate(table.rows[1:4], 1):
                    row_data = [cell.text for cell in row.cells]
                    print(f"     行{j}: {row_data}")
        
        # 查找包含"表"的段落
        print(f"\n3. 表格标题:")
        for para in doc.paragraphs:
            if '表' in para.text and ('数据分析结果' in para.text or '数据说明' in para.text):
                print(f"   - {para.text[:80]}")
        
        print("\n" + "="*80)
        print("验证完成!")
        print("="*80)
        
    except Exception as e:
        print(f"验证失败: {e}")
        import traceback
        traceback.print_exc()

if __name__ == '__main__':
    verify_word_tables()

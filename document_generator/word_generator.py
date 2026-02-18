"""
Word文档生成器
生成结构完整、内容详实的Word分析报告
"""

import json
from typing import Dict, List, Optional, Any
from io import BytesIO
from datetime import datetime
from PIL import Image

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

from analyzer import AnalysisResult, AnalysisDimension


class WordDocumentGenerator:
    """
    Word文档生成器
    生成专业的数据分析报告Word文档
    """
    
    def __init__(self):
        self.doc = None
    
    def generate(self, 
                 analysis_result: AnalysisResult,
                 chart_images: List[str] = None,
                 title: str = "数据分析报告") -> BytesIO:
        """
        生成Word分析报告
        
        Args:
            analysis_result: 分析结果
            chart_images: 图表图像路径列表
            title: 报告标题
        
        Returns:
            Word文档字节流
        """
        self.doc = Document()
        
        # 设置文档样式
        self._setup_styles()
        
        # 1. 封面
        self._add_cover_page(title)
        
        # 2. 执行摘要
        self._add_executive_summary(analysis_result)
        
        # 3. 数据概览
        self._add_data_overview(analysis_result)
        
        # 4. 详细分析
        self._add_detailed_analysis(analysis_result)
        
        # 5. 深度洞察
        self._add_insights(analysis_result, chart_images or [])
        
        # 6. 结论与建议
        self._add_conclusion(analysis_result)
        
        # 7. 附录
        self._add_appendix(analysis_result)
        
        # 保存到内存
        output = BytesIO()
        self.doc.save(output)
        output.seek(0)
        
        return output
    
    def _setup_styles(self):
        """设置文档样式"""
        # 标题1样式
        style = self.doc.styles['Heading 1']
        style.font.name = 'Arial'
        style.font.size = Pt(18)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 51, 102)
        
        # 标题2样式
        style = self.doc.styles['Heading 2']
        style.font.name = 'Arial'
        style.font.size = Pt(14)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 102, 204)
        
        # 正文样式
        style = self.doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)
    
    def _add_cover_page(self, title: str):
        """添加封面"""
        # 标题
        heading = self.doc.add_heading(title, 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 副标题
        subtitle = self.doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run("基于AI的数据分析报告")
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(100, 100, 100)
        
        # 日期
        date_para = self.doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = date_para.add_run(f"生成日期: {datetime.now().strftime('%Y年%m月%d日')}")
        run.font.size = Pt(12)
        
        # 分页
        self.doc.add_page_break()
    
    def _add_executive_summary(self, result: AnalysisResult):
        """添加执行摘要"""
        self.doc.add_heading("执行摘要", level=1)
        
        # 关键发现
        self.doc.add_heading("关键发现", level=2)
        for i, insight in enumerate(result.insights[:5], 1):
            p = self.doc.add_paragraph(style='List Bullet')
            p.add_run(f"{insight.get('title', f'发现{i}')}: ").bold = True
            p.add_run(insight.get('description', '')[:100] + "...")
        
        # 核心结论
        self.doc.add_heading("核心结论", level=2)
        conclusion_para = self.doc.add_paragraph()
        if result.insights:
            conclusion_para.add_run(
                f"基于对{len(result.merged_dimensions)}个分析维度的深入研究，"
                f"我们发现了{len(result.insights)}个关键洞察。"
            )
        
        # 建议行动
        self.doc.add_heading("建议行动", level=2)
        for insight in result.insights[:3]:
            p = self.doc.add_paragraph(style='List Bullet')
            p.add_run(insight.get('action', '继续监控'))
    
    def _add_data_overview(self, result: AnalysisResult):
        """添加数据概览"""
        self.doc.add_heading("数据概览", level=1)
        
        preprocessing = result.preprocessing
        
        # 数据来源
        self.doc.add_heading("数据来源", level=2)
        self.doc.add_paragraph("本报告基于上传的数据文件进行分析。")
        
        # 数据规模
        self.doc.add_heading("数据规模", level=2)
        structure = preprocessing.get('structure_understanding', {})
        columns = structure.get('columns', [])
        
        p = self.doc.add_paragraph()
        p.add_run(f"• 总列数: {len(columns)}\n")
        p.add_run(f"• 分析维度: {len(result.merged_dimensions)}")
        
        # 数据质量
        self.doc.add_heading("数据质量评估", level=2)
        quality = preprocessing.get('quality_assessment', {})
        missing = quality.get('missing_values', {})
        duplicates = quality.get('duplicates', {})
        
        p = self.doc.add_paragraph()
        p.add_run(f"• 缺失值: {len(missing)}列存在缺失\n")
        p.add_run(f"• 重复数据: {duplicates.get('count', 0)}行")
        
        # 分析维度
        self.doc.add_heading("分析维度", level=2)
        p = self.doc.add_paragraph("本报告使用以下分析维度：")
        
        for dim in result.merged_dimensions:
            p = self.doc.add_paragraph(style='List Bullet')
            source_label = {"template": "【模板】", "ai": "【AI】", "user": "【用户】"}.get(dim.source, "")
            p.add_run(f"{source_label} {dim.name}")
    
    def _add_detailed_analysis(self, result: AnalysisResult):
        """添加详细分析"""
        self.doc.add_heading("详细分析", level=1)
        
        eda = result.eda
        
        # 单变量分析
        self.doc.add_heading("1. 单变量分析", level=2)
        univariate = eda.get('univariate_analysis', {})
        
        for col, stats in list(univariate.items())[:3]:
            self.doc.add_heading(f"{col}分布特征", level=3)
            
            p = self.doc.add_paragraph()
            p.add_run(f"均值: {stats.get('mean', 'N/A'):.2f}  ")
            p.add_run(f"中位数: {stats.get('median', 'N/A'):.2f}  ")
            p.add_run(f"标准差: {stats.get('std', 'N/A'):.2f}\n")
            p.add_run(f"最小值: {stats.get('min', 'N/A'):.2f}  ")
            p.add_run(f"最大值: {stats.get('max', 'N/A'):.2f}\n")
            p.add_run(f"偏度: {stats.get('skewness', 'N/A'):.2f}  ")
            p.add_run(f"峰度: {stats.get('kurtosis', 'N/A'):.2f}")
        
        # 双变量分析
        self.doc.add_heading("2. 双变量分析", level=2)
        bivariate = eda.get('bivariate_analysis', {})
        strong_corrs = bivariate.get('strong_correlations', [])
        
        if strong_corrs:
            self.doc.add_paragraph("发现以下强相关关系：")
            for corr in strong_corrs[:3]:
                p = self.doc.add_paragraph(style='List Bullet')
                p.add_run(f"{corr['column1']} vs {corr['column2']}: ")
                p.add_run(f"r = {corr['correlation']:.3f}").bold = True
        else:
            self.doc.add_paragraph("未发现显著的强相关关系。")
        
        # 多变量分析
        self.doc.add_heading("3. 多变量分析", level=2)
        multivariate = eda.get('multivariate_analysis', {})
        
        pca = multivariate.get('pca', {})
        if pca:
            p = self.doc.add_paragraph()
            p.add_run(f"主成分分析（PCA）:\n")
            p.add_run(f"• 成分数: {pca.get('n_components', 'N/A')}\n")
            variance_ratio = pca.get('explained_variance_ratio', [])
            if variance_ratio:
                p.add_run(f"• 方差解释比例: {', '.join([f'{v:.2%}' for v in variance_ratio])}")
        
        recommendations = multivariate.get('recommendations', [])
        if recommendations:
            self.doc.add_paragraph("分析建议：")
            for rec in recommendations:
                p = self.doc.add_paragraph(style='List Bullet')
                p.add_run(rec)
    
    def _add_insights(self, result: AnalysisResult, chart_images: List[str]):
        """添加深度洞察"""
        self.doc.add_heading("深度洞察", level=1)
        
        for i, insight in enumerate(result.insights, 1):
            self.doc.add_heading(f"洞察 {i}: {insight.get('title', '')}", level=2)
            
            # 详细描述
            self.doc.add_heading("描述", level=3)
            self.doc.add_paragraph(insight.get('description', ''))
            
            # 关键发现
            self.doc.add_heading("关键发现", level=3)
            findings = insight.get('key_findings', [])
            for finding in findings:
                p = self.doc.add_paragraph(style='List Bullet')
                p.add_run(finding)
            
            # 数据证据
            self.doc.add_heading("数据证据", level=3)
            evidence = insight.get('data_evidence', [])
            for ev in evidence:
                p = self.doc.add_paragraph(style='List Bullet')
                p.add_run(ev)
            
            # 业务影响
            self.doc.add_heading("业务影响", level=3)
            self.doc.add_paragraph(insight.get('business_impact', '待评估'))
            
            # 行动建议
            self.doc.add_heading("行动建议", level=3)
            p = self.doc.add_paragraph()
            p.add_run(insight.get('action', '继续监控')).bold = True
            
            # 置信度和优先级
            p = self.doc.add_paragraph()
            p.add_run(f"置信度: {insight.get('confidence', '中')}  |  ")
            p.add_run(f"优先级: {insight.get('priority', '中')}")
            
            # 维度来源
            source = insight.get('dimension_source', 'ai')
            source_label = {"template": "模板维度", "ai": "AI维度", "user": "用户维度"}.get(source, "AI维度")
            p = self.doc.add_paragraph()
            p.add_run(f"维度来源: {source_label}").italic = True
            
            # 插入图表（如果有）
            if i <= len(chart_images):
                try:
                    self.doc.add_picture(chart_images[i-1], width=Inches(5.5))
                    last_paragraph = self.doc.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception:
                    pass
            
            # 分页（除了最后一个洞察）
            if i < len(result.insights):
                self.doc.add_page_break()
    
    def _add_conclusion(self, result: AnalysisResult):
        """添加结论与建议"""
        self.doc.add_heading("结论与建议", level=1)
        
        # 总结
        self.doc.add_heading("总结", level=2)
        summary = self.doc.add_paragraph()
        summary.add_run(
            f"本报告通过{len(result.merged_dimensions)}个分析维度，"
            f"对数据进行了全面的探索性分析。共发现{len(result.insights)}个关键洞察，"
            f"涵盖单变量、双变量和多变量分析。"
        )
        
        # 后续行动建议
        self.doc.add_heading("后续行动建议", level=2)
        
        high_priority_insights = [i for i in result.insights if i.get('priority') == '高']
        if high_priority_insights:
            self.doc.add_paragraph("高优先级行动：")
            for insight in high_priority_insights:
                p = self.doc.add_paragraph(style='List Bullet')
                p.add_run(f"{insight.get('title', '')}: ")
                p.add_run(insight.get('action', ''))
        
        # 一般建议
        self.doc.add_paragraph("一般建议：")
        suggestions = [
            "持续监控关键指标的变化趋势",
            "深入分析强相关变量间的因果关系",
            "定期进行数据质量检查",
            "根据业务需求调整分析维度"
        ]
        for suggestion in suggestions:
            p = self.doc.add_paragraph(style='List Bullet')
            p.add_run(suggestion)
    
    def _add_appendix(self, result: AnalysisResult):
        """添加附录"""
        self.doc.add_page_break()
        self.doc.add_heading("附录", level=1)
        
        # 分析方法说明
        self.doc.add_heading("分析方法说明", level=2)
        p = self.doc.add_paragraph()
        p.add_run("本报告采用以下分析方法：\n")
        p.add_run("• 单变量分析: 描述统计、分布特征\n")
        p.add_run("• 双变量分析: 相关性分析、协方差分析\n")
        p.add_run("• 多变量分析: 主成分分析(PCA)")
        
        # 维度来源说明
        self.doc.add_heading("维度来源说明", level=2)
        
        template_dims = [d for d in result.merged_dimensions if d.source == 'template']
        ai_dims = [d for d in result.merged_dimensions if d.source == 'ai']
        user_dims = [d for d in result.merged_dimensions if d.source == 'user']
        
        p = self.doc.add_paragraph()
        p.add_run(f"• 模板维度: {len(template_dims)}个\n")
        p.add_run(f"• AI自动维度: {len(ai_dims)}个\n")
        p.add_run(f"• 用户自定义维度: {len(user_dims)}个")
        
        # 处理日志
        self.doc.add_heading("处理日志", level=2)
        for log in result.processing_log[-10:]:  # 只显示最后10条
            p = self.doc.add_paragraph()
            p.add_run(log).font.size = Pt(9)
            p.add_run("\n")

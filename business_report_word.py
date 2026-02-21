"""
Word报告生成器 - 确保所有表格正确呈现
"""
import io
from datetime import datetime
from typing import List

import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT

from business_semantic_report_v3 import (
    BusinessSemanticReportGeneratorV3,
    BusinessMetric,
    DimensionAnalysis,
    KeyInsight
)


class WordReportGenerator:
    """Word报告生成器 - 完整的表格呈现"""
    
    def __init__(self, generator: BusinessSemanticReportGeneratorV3):
        self.generator = generator
        self.doc = Document()
        self._setup_document()
    
    def _setup_document(self):
        """设置文档样式"""
        # 设置默认字体
        style = self.doc.styles['Normal']
        style.font.name = 'Microsoft YaHei'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        style.font.size = Pt(10.5)
    
    def _set_cell_font(self, cell, text: str, bold: bool = False, size: int = 10):
        """设置单元格字体"""
        cell.text = text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Microsoft YaHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                run.font.size = Pt(size)
                run.font.bold = bold
    
    def _add_heading(self, text: str, level: int = 1):
        """添加标题"""
        heading = self.doc.add_heading(text, level=level)
        for run in heading.runs:
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            if level == 0:
                run.font.size = Pt(18)
            elif level == 1:
                run.font.size = Pt(16)
            elif level == 2:
                run.font.size = Pt(14)
            else:
                run.font.size = Pt(12)
            run.font.bold = True
        return heading
    
    def _add_paragraph(self, text: str, bold: bool = False, size: int = 10):
        """添加段落"""
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        run.font.size = Pt(size)
        run.font.bold = bold
        return p
    
    def _add_dataframe_table(self, df: pd.DataFrame, title: str = ""):
        """添加DataFrame表格到Word"""
        if title:
            self._add_paragraph(title, bold=True, size=11)
            self.doc.add_paragraph()
        
        if df.empty:
            self._add_paragraph("(无数据)")
            return
        
        # 创建表格
        table = self.doc.add_table(rows=1, cols=len(df.columns))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 设置表头
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(df.columns):
            self._set_cell_font(hdr_cells[i], str(col), bold=True, size=10)
            # 设置表头背景色
            hdr_cells[i]._element.get_or_add_tcPr().append(
                parse_xml(r'<w:shd {} w:fill="4472C4"/>'.format(nsdecls('w')))
            )
        
        # 添加数据行
        for _, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                self._set_cell_font(row_cells[i], str(value), size=10)
        
        self.doc.add_paragraph()
    
    def generate_full_report(self, title: str = "业务分析报告") -> io.BytesIO:
        """生成完整的Word报告"""
        # 执行分析
        self.generator._analyze_business_metrics()
        self.generator._analyze_by_dimensions()
        self.generator._extract_key_insights()
        self.generator._generate_recommendations()
        
        # 标题
        title_heading = self._add_heading(title, level=0)
        title_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 元信息
        self._add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", size=10)
        self.doc.add_paragraph()
        
        # 章节1：核心业务指标
        self._add_heading("一、核心业务指标", level=1)
        self._add_paragraph("本章节展示了数据中的核心业务指标汇总统计，反映整体业务规模。")
        self.doc.add_paragraph()
        
        if self.generator.business_metrics:
            # 业务指标表格
            metrics_df = self._format_metrics_table(self.generator.business_metrics)
            self._add_dataframe_table(metrics_df, "核心业务指标统计表")
            
            # 指标说明
            self._add_paragraph("指标说明：", bold=True)
            self._add_paragraph("• 汇总：该指标所有数据的总和，反映整体规模")
            self._add_paragraph("• 平均：该指标的平均值，反映一般水平")
            self._add_paragraph("• 最大/最小：该指标的极值，反映波动范围")
        else:
            self._add_paragraph("未识别到数值型业务指标")
        
        self.doc.add_page_break()
        
        # 章节2：维度分析
        self._add_heading("二、维度分析", level=1)
        self._add_paragraph("本章节按不同维度对业务指标进行分析，识别各维度的表现差异和业务重点。")
        self.doc.add_paragraph()
        
        if self.generator.dimension_analyses:
            for i, analysis in enumerate(self.generator.dimension_analyses[:5], 1):
                self._add_heading(f"2.{i} {analysis.dimension_name}维度 - {analysis.metric_name}", level=2)
                
                # 分析总结
                self._add_paragraph(f"分析总结：{analysis.summary}", bold=True)
                self.doc.add_paragraph()
                
                # 维度数据表格
                dim_df = self._format_dimension_table(analysis)
                self._add_dataframe_table(dim_df, f"{analysis.dimension_name}维度详细数据")
                
                # 详细说明
                if analysis.top_values:
                    top1 = analysis.top_values[0]
                    self._add_paragraph("详细说明：", bold=True)
                    self._add_paragraph(f"• 表现最优：{top1.get(analysis.dimension_name, '未知')}，汇总值{top1.get('汇总', 0):.2f}")
                    if len(analysis.top_values) > 1:
                        top2 = analysis.top_values[1]
                        self._add_paragraph(f"• 排名第二：{top2.get(analysis.dimension_name, '未知')}，汇总值{top2.get('汇总', 0):.2f}")
                    self._add_paragraph(f"• 数据覆盖：共{len(analysis.data)}个类别，平均每个类别{analysis.data['平均'].mean():.2f}")
                
                self.doc.add_paragraph()
        else:
            self._add_paragraph("未识别到有效的维度分析")
        
        self.doc.add_page_break()
        
        # 章节3：关键洞察
        self._add_heading("三、关键洞察", level=1)
        self._add_paragraph("本章节基于数据分析提取关键业务洞察，每个洞察都有完整的数据表格支撑。")
        self.doc.add_paragraph()
        
        if self.generator.key_insights:
            for i, insight in enumerate(self.generator.key_insights, 1):
                type_names = {
                    'anomaly': '⚠️ 异常发现',
                    'opportunity': '💡 业务机会',
                    'pattern': '📊 数据模式',
                    'risk': '⚡ 潜在风险'
                }
                type_name = type_names.get(insight.insight_type, insight.insight_type)
                
                self._add_heading(f"3.{i} {insight.title}", level=2)
                
                self._add_paragraph(f"洞察类型：{type_name}", bold=True)
                self.doc.add_paragraph()
                
                self._add_paragraph(f"洞察描述：{insight.description}")
                self.doc.add_paragraph()
                
                # 支撑数据表格
                if insight.supporting_table is not None and not insight.supporting_table.empty:
                    self._add_dataframe_table(insight.supporting_table, "支撑数据详情")
                elif insight.supporting_data:
                    support_df = pd.DataFrame([insight.supporting_data])
                    self._add_dataframe_table(support_df, "支撑数据详情")
                
                self._add_paragraph(f"置信度：{insight.confidence} | 优先级：{insight.priority}", size=9)
                self.doc.add_paragraph()
        else:
            self._add_paragraph("暂无关键洞察")
        
        self.doc.add_page_break()
        
        # 章节4：业务建议
        self._add_heading("四、业务建议", level=1)
        self._add_paragraph("本章节基于关键洞察提出可执行的业务建议。")
        self.doc.add_paragraph()
        
        if self.generator.recommendations:
            for i, rec in enumerate(self.generator.recommendations, 1):
                type_names = {
                    'improvement': '🔧 改进建议',
                    'opportunity': '🚀 机会建议',
                    'optimization': '⚙️ 优化建议'
                }
                type_name = type_names.get(rec.recommendation_type, rec.recommendation_type)
                
                self._add_heading(f"4.{i} {rec.title}", level=2)
                
                self._add_paragraph(f"建议类型：{type_name}", bold=True)
                self.doc.add_paragraph()
                
                self._add_paragraph(f"建议内容：{rec.description}")
                self.doc.add_paragraph()
                
                if rec.expected_impact:
                    self._add_paragraph(f"预期效果：{rec.expected_impact}")
                    self.doc.add_paragraph()
                
                self._add_paragraph(f"实施难度：{rec.implementation_difficulty}", size=9)
                self.doc.add_paragraph()
        else:
            self._add_paragraph("暂无业务建议")
        
        # 保存文档
        output = io.BytesIO()
        self.doc.save(output)
        output.seek(0)
        return output
    
    def _format_metrics_table(self, metrics: List[BusinessMetric]) -> pd.DataFrame:
        """格式化业务指标表"""
        data = []
        for m in metrics:
            data.append({
                '指标名称': m.name,
                '指标值': m.format_str.format(m.value),
                '单位': m.unit,
                '指标说明': m.description
            })
        return pd.DataFrame(data)
    
    def _format_dimension_table(self, analysis: DimensionAnalysis) -> pd.DataFrame:
        """格式化维度分析表"""
        df = analysis.data.copy()
        # 添加排名列
        df['排名'] = range(1, len(df) + 1)
        # 重新排列列顺序
        cols = ['排名', analysis.dimension_name, '汇总', '平均', '计数']
        df = df[[c for c in cols if c in df.columns]]
        return df


# 导入需要的模块
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls


def generate_complete_word_report(df: pd.DataFrame, title: str = "业务分析报告") -> bytes:
    """生成完整的Word报告"""
    # 创建分析器
    analyzer = BusinessSemanticReportGeneratorV3(df)
    
    # 创建Word生成器
    word_gen = WordReportGenerator(analyzer)
    
    # 生成报告
    output = word_gen.generate_full_report(title)
    
    return output.getvalue()

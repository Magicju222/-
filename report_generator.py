"""
Summary Report Generator
生成数据分析报告的Markdown和Word文档
与网页显示保持一致
"""

import io
import json
import re
from typing import Any, Dict, List, Optional, Union
from datetime import datetime

import pandas as pd
import numpy as np
from scipy import stats
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


def _set_chinese_font(run, font_name='Microsoft YaHei', font_size=10.5, bold=False, italic=False):
    """设置中文字体"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def _extract_insights_from_agent(agent_result) -> List[Dict]:
    """从Agent结果中提取洞察"""
    insights = []
    
    if hasattr(agent_result, 'insights') and agent_result.insights:
        for insight in agent_result.insights:
            insights.append({
                'title': insight.get('title', ''),
                'description': insight.get('description', ''),
                'key_findings': insight.get('key_findings', []),
                'action': insight.get('action', ''),
                'confidence': insight.get('confidence', '中'),
                'priority': insight.get('priority', '中')
            })
    
    return insights


def _extract_observations(agent_result) -> List[Dict]:
    """提取分析步骤（包含完整的步骤信息和所有表格数据）"""
    observations = []
    
    if not hasattr(agent_result, 'steps') or not agent_result.steps:
        return observations
    
    for step in agent_result.steps:
        observation = getattr(step, 'observation', None)
        if observation and observation.strip():
            # 提取所有表格数据
            tables = _extract_all_tables_from_observation(observation)
            
            # 提取非表格的文本内容
            text_content = _extract_text_without_tables(observation)
            
            # 如果没有表格，尝试从文本中提取关键发现并转换为表格
            if not tables and text_content:
                extracted_table = _extract_key_findings_to_table(text_content)
                if extracted_table:
                    tables.append(extracted_table)
            
            observations.append({
                'step_number': getattr(step, 'step_number', 0),
                'action': getattr(step, 'action', 'unknown'),
                'thought': getattr(step, 'thought', ''),
                'observation': observation,
                'text_content': text_content,
                'tables': tables,
                'has_table': len(tables) > 0,
                'table_count': len(tables),
                'tool_result': getattr(step, 'tool_result', None)
            })
    
    return observations


def _extract_text_without_tables(observation: str) -> str:
    """提取观察结果中的非表格文本内容"""
    lines = observation.strip().split('\n')
    text_lines = []
    in_table = False
    
    for line in lines:
        stripped = line.strip()
        
        # 检测是否在表格中
        if '|' in stripped or '\t' in stripped:
            in_table = True
            continue
        else:
            if in_table:
                in_table = False
            if stripped:
                text_lines.append(stripped)
    
    return '\n'.join(text_lines)


def _parse_observation_to_table(observation: str) -> Optional[Dict]:
    """从观察结果文本中解析表格数据 - 增强版"""
    lines = observation.strip().split('\n')
    
    # 查找表格行（包含 | 或制表符分隔的数据）
    table_lines = []
    for line in lines:
        line = line.strip()
        if '|' in line or '\t' in line:
            table_lines.append(line)
    
    if len(table_lines) < 2:
        return None
    
    # 解析表格
    headers = []
    rows = []
    
    for i, line in enumerate(table_lines):
        # 分割单元格
        if '|' in line:
            cells = [cell.strip() for cell in line.split('|')]
            # 过滤空单元格，但保留有意义的空值
            cells = [c for c in cells if c or c == '']
        else:
            cells = [cell.strip() for cell in line.split('\t')]
        
        if not cells or all(c == '' for c in cells):
            continue
        
        # 跳过Markdown分隔符行
        if all(c.replace('-', '').replace(':', '').strip() == '' for c in cells):
            continue
        
        if i == 0 or not headers:
            headers = cells
        else:
            rows.append(cells)
    
    if headers and rows:
        return {'headers': headers, 'rows': rows}
    
    return None


def _extract_all_tables_from_observation(observation: str) -> List[Dict]:
    """从观察结果中提取所有表格数据"""
    tables = []
    lines = observation.strip().split('\n')
    
    current_table_lines = []
    in_table = False
    
    for line in lines:
        stripped = line.strip()
        
        # 检测表格行
        if '|' in stripped or '\t' in stripped:
            if not in_table:
                in_table = True
                current_table_lines = []
            current_table_lines.append(stripped)
        else:
            if in_table and current_table_lines:
                # 解析当前表格
                table = _parse_table_lines(current_table_lines)
                if table:
                    tables.append(table)
                current_table_lines = []
                in_table = False
    
    # 处理最后一个表格
    if in_table and current_table_lines:
        table = _parse_table_lines(current_table_lines)
        if table:
            tables.append(table)
    
    return tables


def _parse_table_lines(table_lines: List[str]) -> Optional[Dict]:
    """解析表格行数据"""
    if len(table_lines) < 2:
        return None
    
    headers = []
    rows = []
    
    for i, line in enumerate(table_lines):
        # 分割单元格
        if '|' in line:
            cells = [cell.strip() for cell in line.split('|')]
            cells = [c for c in cells if c or c == '']
        else:
            cells = [cell.strip() for cell in line.split('\t')]
        
        if not cells or all(c == '' for c in cells):
            continue
        
        # 跳过Markdown分隔符行
        if all(c.replace('-', '').replace(':', '').strip() == '' for c in cells):
            continue
        
        if i == 0 or not headers:
            headers = cells
        else:
            # 确保每行单元格数与表头一致
            if len(cells) < len(headers):
                cells.extend([''] * (len(headers) - len(cells)))
            elif len(cells) > len(headers):
                cells = cells[:len(headers)]
            rows.append(cells)
    
    if headers and rows:
        return {'headers': headers, 'rows': rows}
    
    return None


def _extract_key_findings_to_table(text_content: str) -> Optional[Dict]:
    """
    从文本中提取关键发现并转换为表格
    例如：从 "1. 华东地区表现突出：完成率102.5%，同比增长18%" 提取为表格行
    """
    import re
    
    lines = text_content.strip().split('\n')
    rows = []
    
    # 匹配模式：序号. 地区/指标：数值
    patterns = [
        r'(\d+)\.\s*([^：:]+)[：:]\s*(.+)',  # 1. 华东地区：完成率102.5%
        r'-\s*([^：:]+)[：:]\s*(.+)',        # - 华东地区：完成率102.5%
        r'•\s*([^：:]+)[：:]\s*(.+)',        # • 华东地区：完成率102.5%
    ]
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        for pattern in patterns:
            match = re.match(pattern, line)
            if match:
                if len(match.groups()) == 3:
                    # 有序号
                    _, name, value = match.groups()
                    rows.append([name.strip(), value.strip()])
                else:
                    # 无序号
                    name, value = match.groups()
                    rows.append([name.strip(), value.strip()])
                break
    
    if rows:
        return {
            'headers': ['项目', '数据/描述'],
            'rows': rows
        }
    
    return None


def _extract_final_report_content(agent_result) -> str:
    """提取最终报告内容"""
    if hasattr(agent_result, 'final_report') and agent_result.final_report:
        return agent_result.final_report
    return ""


def generate_markdown_report(agent_result, dfs, sheet_name) -> str:
    """生成Markdown报告 - 与网页显示一致"""
    
    # 获取洞察
    insights = _extract_insights_from_agent(agent_result)
    
    # 获取观察结果
    observations = _extract_observations(agent_result)
    
    # 获取最终报告
    final_report = _extract_final_report_content(agent_result)
    
    report_lines = []
    
    # 报告标题
    report_lines.append(f"# 数据分析报告")
    report_lines.append("")
    report_lines.append(f"**分析对象**: {sheet_name}")
    report_lines.append(f"**生成时间**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    report_lines.append("")
    report_lines.append("---")
    report_lines.append("")
    
    # 一、分析维度
    report_lines.append("## 一、分析维度")
    report_lines.append("")
    
    if hasattr(agent_result, 'merged_dimensions') and agent_result.merged_dimensions:
        template_dims = [d for d in agent_result.merged_dimensions if d.source == 'template']
        ai_dims = [d for d in agent_result.merged_dimensions if d.source == 'ai']
        user_dims = [d for d in agent_result.merged_dimensions if d.source == 'user']
        
        if template_dims:
            report_lines.append("**模板维度**:")
            for dim in template_dims[:5]:
                report_lines.append(f"- {dim.name}")
            report_lines.append("")
        
        if ai_dims:
            report_lines.append("**AI自动维度**:")
            for dim in ai_dims[:5]:
                report_lines.append(f"- {dim.name}")
            report_lines.append("")
        
        if user_dims:
            report_lines.append("**用户维度**:")
            for dim in user_dims[:5]:
                report_lines.append(f"- {dim.name}")
            report_lines.append("")
    
    report_lines.append("---")
    report_lines.append("")
    
    # 二、关键洞察（使用网页版的洞察内容）
    report_lines.append("## 二、关键洞察")
    report_lines.append("")
    
    if insights:
        for i, insight in enumerate(insights[:5], 1):
            report_lines.append(f"### 洞察 {i}: {insight['title']}")
            report_lines.append("")
            
            if insight['description']:
                report_lines.append(insight['description'])
                report_lines.append("")
            
            if insight['key_findings']:
                report_lines.append("**关键发现**:")
                for finding in insight['key_findings']:
                    report_lines.append(f"- {finding}")
                report_lines.append("")
            
            if insight['action']:
                report_lines.append(f"**行动建议**: {insight['action']}")
                report_lines.append("")
            
            report_lines.append(f"*置信度: {insight['confidence']} | 优先级: {insight['priority']}*")
            report_lines.append("")
    else:
        report_lines.append("暂无关键洞察。")
        report_lines.append("")
    
    report_lines.append("---")
    report_lines.append("")
    
    # 三、详细分析结果
    report_lines.append("## 三、详细分析结果")
    report_lines.append("")
    
    if observations:
        for i, obs in enumerate(observations, 1):
            # 只显示数据表格和分析说明，不显示步骤编号、执行动作、分析思路
            
            # 显示所有表格数据（完整数据，不限制行数）
            if obs['has_table'] and obs['tables']:
                for table_idx, table_data in enumerate(obs['tables'], 1):
                    if table_data['headers'] and table_data['rows']:
                        # 表格标题
                        if obs['table_count'] > 1:
                            report_lines.append(f"**数据表 {table_idx}：**")
                            report_lines.append("")
                        
                        # 生成Markdown表格
                        # 表头
                        header_line = "| " + " | ".join(table_data['headers']) + " |"
                        report_lines.append(header_line)
                        
                        # 分隔符
                        separator = "|" + "|".join(["---" for _ in table_data['headers']]) + "|"
                        report_lines.append(separator)
                        
                        # 数据行（显示完整数据，不限制行数）
                        for row in table_data['rows']:
                            row_line = "| " + " | ".join([str(cell) for cell in row]) + " |"
                            report_lines.append(row_line)
                        
                        report_lines.append("")
                        
                        # 添加文字说明
                        total_rows = len(table_data['rows'])
                        report_lines.append(f"*上表展示了{total_rows}条记录的分析结果，包含{len(table_data['headers'])}个维度。*")
                        report_lines.append("")
            
            # 显示非表格的文本内容（分析说明）
            if obs.get('text_content'):
                report_lines.append("**分析说明：**")
                report_lines.append(obs['text_content'])
                report_lines.append("")
            
            # 分隔不同步骤的结果
            if i < len(observations):
                report_lines.append("---")
                report_lines.append("")
    else:
        report_lines.append("暂无详细分析结果。")
        report_lines.append("")
    
    report_lines.append("---")
    report_lines.append("")
    
    # 四、最终报告摘要
    if final_report:
        report_lines.append("## 四、分析摘要")
        report_lines.append("")
        report_lines.append(final_report)
        report_lines.append("")
        report_lines.append("---")
        report_lines.append("")
    
    # 五、结论与建议
    report_lines.append("## 五、结论与建议")
    report_lines.append("")
    
    if insights:
        report_lines.append("基于上述分析，我们得出以下结论和建议：")
        report_lines.append("")
        
        for i, insight in enumerate(insights[:3], 1):
            report_lines.append(f"{i}. **{insight['title']}**")
            if insight['key_findings']:
                report_lines.append(f"   - 关键发现: {', '.join(insight['key_findings'][:2])}")
            if insight['action']:
                report_lines.append(f"   - 建议措施: {insight['action']}")
            report_lines.append("")
    else:
        report_lines.append("建议持续监控关键业务指标的变化趋势，定期评估业务运营状况。")
        report_lines.append("")
    
    report_lines.append("")
    report_lines.append(f"*报告生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}*")
    
    return "\n".join(report_lines)


def generate_word_report(agent_result, dfs, sheet_name) -> io.BytesIO:
    """生成Word报告 - 与网页显示一致"""
    
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    # 获取数据
    insights = _extract_insights_from_agent(agent_result)
    observations = _extract_observations(agent_result)
    final_report = _extract_final_report_content(agent_result)
    
    # 标题
    title = doc.add_heading('数据分析报告', 0)
    for run in title.runs:
        _set_chinese_font(run, font_size=18, bold=True)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 元信息
    p = doc.add_paragraph()
    run = p.add_run(f"分析对象: {sheet_name}")
    _set_chinese_font(run)
    p = doc.add_paragraph()
    run = p.add_run(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    _set_chinese_font(run)
    
    doc.add_paragraph()
    
    # 一、分析维度
    heading1 = doc.add_heading('一、分析维度', level=1)
    for run in heading1.runs:
        _set_chinese_font(run, font_size=14, bold=True)
    
    if hasattr(agent_result, 'merged_dimensions') and agent_result.merged_dimensions:
        template_dims = [d for d in agent_result.merged_dimensions if d.source == 'template']
        ai_dims = [d for d in agent_result.merged_dimensions if d.source == 'ai']
        user_dims = [d for d in agent_result.merged_dimensions if d.source == 'user']
        
        if template_dims:
            p = doc.add_paragraph()
            run = p.add_run("模板维度:")
            _set_chinese_font(run, bold=True)
            for dim in template_dims[:5]:
                p = doc.add_paragraph(f"• {dim.name}", style='List Bullet')
                for run in p.runs:
                    _set_chinese_font(run)
        
        if ai_dims:
            p = doc.add_paragraph()
            run = p.add_run("AI自动维度:")
            _set_chinese_font(run, bold=True)
            for dim in ai_dims[:5]:
                p = doc.add_paragraph(f"• {dim.name}", style='List Bullet')
                for run in p.runs:
                    _set_chinese_font(run)
        
        if user_dims:
            p = doc.add_paragraph()
            run = p.add_run("用户维度:")
            _set_chinese_font(run, bold=True)
            for dim in user_dims[:5]:
                p = doc.add_paragraph(f"• {dim.name}", style='List Bullet')
                for run in p.runs:
                    _set_chinese_font(run)
    
    doc.add_page_break()
    
    # 二、关键洞察
    heading2 = doc.add_heading('二、关键洞察', level=1)
    for run in heading2.runs:
        _set_chinese_font(run, font_size=14, bold=True)
    
    if insights:
        for i, insight in enumerate(insights[:5], 1):
            insight_heading = doc.add_heading(f"洞察 {i}: {insight['title']}", level=2)
            for run in insight_heading.runs:
                _set_chinese_font(run, font_size=12, bold=True)
            
            if insight['description']:
                p = doc.add_paragraph(insight['description'])
                for run in p.runs:
                    _set_chinese_font(run)
            
            if insight['key_findings']:
                p = doc.add_paragraph()
                run = p.add_run("关键发现:")
                _set_chinese_font(run, bold=True)
                for finding in insight['key_findings']:
                    p = doc.add_paragraph(f"• {finding}", style='List Bullet')
                    for run in p.runs:
                        _set_chinese_font(run)
            
            if insight['action']:
                p = doc.add_paragraph()
                run = p.add_run(f"行动建议: {insight['action']}")
                _set_chinese_font(run)
            
            p = doc.add_paragraph()
            run = p.add_run(f"置信度: {insight['confidence']} | 优先级: {insight['priority']}")
            _set_chinese_font(run, font_size=9)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            doc.add_paragraph()
    else:
        p = doc.add_paragraph("暂无关键洞察。")
        for run in p.runs:
            _set_chinese_font(run)
    
    doc.add_page_break()
    
    # 三、详细分析结果
    heading3 = doc.add_heading('三、详细分析结果', level=1)
    for run in heading3.runs:
        _set_chinese_font(run, font_size=14, bold=True)
    
    # 添加章节说明
    p = doc.add_paragraph()
    run = p.add_run("本章节展示了数据分析的详细结果，所有数据均以规范表格形式呈现，便于查阅和分析。")
    _set_chinese_font(run, font_size=10)
    doc.add_paragraph()
    
    if observations:
        for i, obs in enumerate(observations, 1):
            # 显示所有表格数据（完整数据，不限制行数）
            if obs['has_table'] and obs['tables']:
                for table_idx, table_data in enumerate(obs['tables'], 1):
                    if table_data['headers'] and table_data['rows']:
                        # 表格标题
                        p = doc.add_paragraph()
                        if obs['table_count'] > 1:
                            run = p.add_run(f"表 {i}-{table_idx}：数据分析结果")
                        else:
                            run = p.add_run(f"表 {i}：数据分析结果")
                        _set_chinese_font(run, font_size=11, bold=True)
                        
                        # 表格前说明
                        p = doc.add_paragraph()
                        run = p.add_run("下表展示了分析得到的关键数据指标，可用于进一步的数据分析和决策支持。")
                        _set_chinese_font(run, font_size=9)
                        
                        # 创建表格（显示完整数据，不限制行数）
                        table = doc.add_table(rows=1, cols=len(table_data['headers']))
                        table.style = 'Table Grid'  # 使用更清晰的网格样式
                        
                        # 设置表格自动调整
                        table.autofit = False
                        table.allow_autofit = False
                        
                        # 表头
                        hdr_cells = table.rows[0].cells
                        for j, header in enumerate(table_data['headers']):
                            hdr_cells[j].text = str(header)
                            # 设置表头背景色和字体
                            hdr_cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for paragraph in hdr_cells[j].paragraphs:
                                for run in paragraph.runs:
                                    _set_chinese_font(run, bold=True, font_size=10)
                                    run.font.color.rgb = RGBColor(255, 255, 255)  # 白色字体
                            # 设置背景色（深蓝色）
                            shading_elm = parse_xml(r'<w:shd {} w:fill="1F4E79"/>'.format(nsdecls('w')))
                            hdr_cells[j]._tc.get_or_add_tcPr().append(shading_elm)
                        
                        # 数据行（显示完整数据）
                        for row_idx, row in enumerate(table_data['rows']):
                            row_cells = table.add_row().cells
                            for j, cell in enumerate(row):
                                if j < len(row_cells):
                                    row_cells[j].text = str(cell)
                                    # 居中对齐
                                    row_cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    for paragraph in row_cells[j].paragraphs:
                                        for run in paragraph.runs:
                                            _set_chinese_font(run, font_size=9)
                                    # 交替行背景色
                                    if row_idx % 2 == 1:
                                        shading_elm = parse_xml(r'<w:shd {} w:fill="F2F2F2"/>'.format(nsdecls('w')))
                                        row_cells[j]._tc.get_or_add_tcPr().append(shading_elm)
                        
                        # 表格后说明
                        total_rows = len(table_data['rows'])
                        p = doc.add_paragraph()
                        run = p.add_run(f"【数据说明】上表共包含{total_rows}条记录，{len(table_data['headers'])}个数据维度。这些数据反映了业务分析的关键指标，建议结合实际情况进行深入分析。")
                        _set_chinese_font(run, font_size=9, italic=True)
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        
                        doc.add_paragraph()
            
            # 显示非表格的文本内容（分析说明）
            if obs.get('text_content'):
                p = doc.add_paragraph()
                run = p.add_run("【分析说明】")
                _set_chinese_font(run, bold=True, font_size=10)
                
                p = doc.add_paragraph(obs['text_content'])
                for run in p.runs:
                    _set_chinese_font(run, font_size=9)
                
                doc.add_paragraph()
            
            # 分隔不同步骤的结果
            if i < len(observations):
                doc.add_paragraph()
    else:
        p = doc.add_paragraph("暂无详细分析结果。")
        for run in p.runs:
            _set_chinese_font(run)
    
    doc.add_page_break()
    
    # 四、分析摘要
    if final_report:
        heading4 = doc.add_heading('四、分析摘要', level=1)
        for run in heading4.runs:
            _set_chinese_font(run, font_size=14, bold=True)
        
        p = doc.add_paragraph(final_report)
        for run in p.runs:
            _set_chinese_font(run)
        
        doc.add_page_break()
    
    # 五、结论与建议
    heading5 = doc.add_heading('五、结论与建议', level=1)
    for run in heading5.runs:
        _set_chinese_font(run, font_size=14, bold=True)
    
    if insights:
        p = doc.add_paragraph("基于上述分析，我们得出以下结论和建议：")
        for run in p.runs:
            _set_chinese_font(run)
        
        for i, insight in enumerate(insights[:3], 1):
            p = doc.add_paragraph()
            run = p.add_run(f"{i}. {insight['title']}")
            _set_chinese_font(run, bold=True)
            
            if insight['key_findings']:
                p = doc.add_paragraph(f"关键发现: {', '.join(insight['key_findings'][:2])}")
                for run in p.runs:
                    _set_chinese_font(run)
            
            if insight['action']:
                p = doc.add_paragraph(f"建议措施: {insight['action']}")
                for run in p.runs:
                    _set_chinese_font(run)
    else:
        p = doc.add_paragraph("建议持续监控关键业务指标的变化趋势，定期评估业务运营状况。")
        for run in p.runs:
            _set_chinese_font(run)
    
    doc.add_paragraph()
    
    # 页脚
    p = doc.add_paragraph()
    run = p.add_run(f"报告生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    _set_chinese_font(run, font_size=9)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # 保存到内存
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    return doc_io

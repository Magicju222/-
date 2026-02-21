"""
业务语义分析报告生成器
基于数据自动识别业务指标、维度、趋势，生成专业的业务分析报告
"""

import io
import json
import re
from typing import Any, Dict, List, Optional, Union, Tuple
from datetime import datetime
from dataclasses import dataclass, field

import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from llm_client import LLMClient, LLMConfig


@dataclass
class BusinessMetric:
    """业务指标"""
    name: str
    value: float
    unit: str = ""
    description: str = ""
    aggregation_type: str = ""  # sum, avg, max, min, std, count
    format_str: str = "{:.2f}"


@dataclass
class DimensionAnalysis:
    """维度分析结果"""
    dimension_name: str
    metric_name: str
    data: pd.DataFrame  # 包含维度值和对应指标值
    top_values: List[Dict] = field(default_factory=list)
    bottom_values: List[Dict] = field(default_factory=list)


@dataclass
class TrendAnalysis:
    """趋势分析结果"""
    time_field: str
    metric_name: str
    data: pd.DataFrame  # 包含时间和指标值
    trend_direction: str = ""  # up, down, stable
    growth_rate: float = 0.0
    peak_point: Optional[Dict] = None
    valley_point: Optional[Dict] = None


@dataclass
class KeyInsight:
    """关键洞察"""
    title: str
    description: str
    insight_type: str  # anomaly, opportunity, pattern, risk
    supporting_data: Dict[str, Any]
    confidence: str = "中"
    priority: str = "中"


@dataclass
class Recommendation:
    """业务建议"""
    title: str
    description: str
    recommendation_type: str  # improvement, opportunity, optimization
    expected_impact: str = ""
    implementation_difficulty: str = "中"
    supporting_insights: List[str] = field(default_factory=list)


class BusinessSemanticReportGenerator:
    """
    业务语义分析报告生成器
    自动识别业务指标、维度、趋势，生成专业的业务分析报告
    """
    
    def __init__(self, df: Optional[pd.DataFrame] = None, llm_client: Optional[LLMClient] = None):
        self.llm_client = llm_client
        self.df: Optional[pd.DataFrame] = None
        self.numeric_columns: List[str] = []
        self.categorical_columns: List[str] = []
        self.datetime_columns: List[str] = []
        self.business_metrics: List[BusinessMetric] = []
        self.dimension_analyses: List[DimensionAnalysis] = []
        self.trend_analyses: List[TrendAnalysis] = []
        self.key_insights: List[KeyInsight] = []
        self.recommendations: List[Recommendation] = []
        
        # 如果提供了DataFrame，立即检测列类型
        if df is not None:
            self._detect_column_types(df)
        
    def _detect_column_types(self, df: pd.DataFrame) -> None:
        """检测列类型"""
        self.df = df.copy()
        self.numeric_columns = []
        self.categorical_columns = []
        self.datetime_columns = []
        
        for col in df.columns:
            # 跳过明显是ID的列
            if any(keyword in col.lower() for keyword in ['id', '编号', '序号', 'code']):
                continue
                
            if pd.api.types.is_datetime64_any_dtype(df[col]):
                self.datetime_columns.append(col)
            elif pd.api.types.is_numeric_dtype(df[col]):
                # 检查是否为分类编码（数值范围很小）
                unique_count = df[col].nunique()
                total_count = len(df)
                if unique_count <= 10 and unique_count / total_count < 0.1:
                    self.categorical_columns.append(col)
                else:
                    self.numeric_columns.append(col)
            else:
                # 尝试转换为日期
                try:
                    pd.to_datetime(df[col], errors='raise')
                    self.datetime_columns.append(col)
                except:
                    # 检查是否为分类变量
                    unique_count = df[col].nunique()
                    total_count = len(df)
                    if unique_count <= 50 or unique_count / total_count < 0.3:
                        self.categorical_columns.append(col)
    
    def _get_metric_unit(self, column_name: str) -> str:
        """根据列名推断单位"""
        name_lower = column_name.lower()
        
        # 金额相关
        if any(kw in name_lower for kw in ['金额', '价格', '费用', '成本', '收入', '销售额', '营收', 'price', 'amount', 'cost', 'revenue', 'sales']):
            return "元"
        # 数量相关
        elif any(kw in name_lower for kw in ['数量', '销量', 'count', 'quantity', 'volume', 'qty']):
            return "个/件"
        # 百分比相关
        elif any(kw in name_lower for kw in ['率', '比例', '占比', 'percent', 'rate', 'ratio']):
            return "%"
        # 时间相关
        elif any(kw in name_lower for kw in ['时间', '时长', '周期', 'time', 'duration']):
            return "小时"
        # 人数相关
        elif any(kw in name_lower for kw in ['人数', '员工', '客户', 'people', 'staff', 'customer']):
            return "人"
        else:
            return ""
    
    def _get_metric_description(self, column_name: str, aggregation_type: str) -> str:
        """生成指标描述"""
        descriptions = {
            'sum': f"{column_name}的总和",
            'avg': f"{column_name}的平均值",
            'max': f"{column_name}的最大值",
            'min': f"{column_name}的最小值",
            'std': f"{column_name}的标准差",
            'count': f"{column_name}的计数"
        }
        return descriptions.get(aggregation_type, f"{column_name}的{aggregation_type}")
    
    def _analyze_business_metrics(self) -> List[BusinessMetric]:
        """
        自动识别核心业务指标，计算指标值
        返回规范的业务指标表结构
        """
        metrics = []
        
        if self.df is None or self.df.empty:
            return metrics
        
        for col in self.numeric_columns:
            series = self.df[col].dropna()
            
            if series.empty:
                continue
            
            unit = self._get_metric_unit(col)
            
            # 计算各种聚合指标
            aggregations = [
                ('sum', series.sum(), "{:.2f}"),
                ('avg', series.mean(), "{:.2f}"),
                ('max', series.max(), "{:.2f}"),
                ('min', series.min(), "{:.2f}"),
                ('std', series.std(), "{:.2f}"),
                ('count', series.count(), "{:.0f}")
            ]
            
            # 根据列名特征决定主要展示哪些指标
            name_lower = col.lower()
            
            # 金额类指标：优先展示sum和avg
            if any(kw in name_lower for kw in ['金额', '价格', '费用', '成本', '收入', '销售额', '营收', 'price', 'amount', 'cost', 'revenue', 'sales']):
                priority_aggs = ['sum', 'avg', 'max', 'min']
            # 比率类指标：优先展示avg
            elif any(kw in name_lower for kw in ['率', '比例', '占比', 'percent', 'rate', 'ratio']):
                priority_aggs = ['avg', 'max', 'min', 'std']
            # 其他：展示全部
            else:
                priority_aggs = ['sum', 'avg', 'max', 'min', 'std', 'count']
            
            for agg_type, value, fmt in aggregations:
                if agg_type in priority_aggs:
                    metric = BusinessMetric(
                        name=f"{col} ({self._get_agg_name(agg_type)})",
                        value=float(value),
                        unit=unit,
                        description=self._get_metric_description(col, agg_type),
                        aggregation_type=agg_type,
                        format_str=fmt
                    )
                    metrics.append(metric)
        
        self.business_metrics = metrics
        return metrics
    
    def _get_agg_name(self, agg_type: str) -> str:
        """获取聚合类型中文名"""
        names = {
            'sum': '汇总',
            'avg': '平均',
            'max': '最大',
            'min': '最小',
            'std': '标准差',
            'count': '计数'
        }
        return names.get(agg_type, agg_type)
    
    def _analyze_by_dimensions(self) -> List[DimensionAnalysis]:
        """
        识别数据中的分类维度，按维度统计业务指标
        返回规范的维度分析表结构
        """
        analyses = []
        
        if self.df is None or self.df.empty:
            return analyses
        
        # 对每个分类维度进行分析
        for dim_col in self.categorical_columns:
            # 跳过唯一值过多的维度
            if self.df[dim_col].nunique() > 100:
                continue
            
            # 对每个数值指标进行分析
            for metric_col in self.numeric_columns[:5]:  # 限制指标数量
                try:
                    # 按维度分组计算指标（显式设置 observed=False 避免警告）
                    grouped = self.df.groupby(dim_col, observed=False)[metric_col].agg([
                        ('汇总', 'sum'),
                        ('平均', 'mean'),
                        ('计数', 'count')
                    ]).reset_index()
                    
                    grouped.columns = [dim_col, '汇总', '平均', '计数']
                    
                    # 按汇总值排序
                    grouped = grouped.sort_values('汇总', ascending=False)
                    
                    # 获取Top和Bottom
                    top_values = grouped.head(5).to_dict('records')
                    bottom_values = grouped.tail(5).to_dict('records')
                    
                    analysis = DimensionAnalysis(
                        dimension_name=dim_col,
                        metric_name=metric_col,
                        data=grouped,
                        top_values=top_values,
                        bottom_values=bottom_values
                    )
                    analyses.append(analysis)
                    
                except Exception as e:
                    continue
        
        self.dimension_analyses = analyses
        return analyses
    
    def _analyze_trends(self) -> List[TrendAnalysis]:
        """
        识别时间字段，按时间统计业务指标
        返回规范的趋势分析表结构
        """
        analyses = []
        
        if self.df is None or self.df.empty:
            return analyses
        
        # 处理时间列
        for time_col in self.datetime_columns:
            try:
                # 转换为日期时间
                time_series = pd.to_datetime(self.df[time_col], errors='coerce')
                
                if time_series.isna().all():
                    continue
                
                # 对每个数值指标进行趋势分析
                for metric_col in self.numeric_columns[:3]:  # 限制指标数量
                    try:
                        # 创建临时DataFrame
                        temp_df = pd.DataFrame({
                            'time': time_series,
                            'metric': self.df[metric_col]
                        }).dropna()
                        
                        if temp_df.empty:
                            continue
                        
                        # 按时间聚合（按天）
                        temp_df['date'] = temp_df['time'].dt.date
                        daily_data = temp_df.groupby('date', observed=False)['metric'].sum().reset_index()
                        daily_data.columns = ['日期', '数值']
                        
                        if len(daily_data) < 2:
                            continue
                        
                        # 计算趋势
                        x = np.arange(len(daily_data))
                        y = daily_data['数值'].values
                        
                        # 简单线性回归计算趋势
                        slope = np.polyfit(x, y, 1)[0]
                        
                        # 判断趋势方向
                        if slope > 0.01 * np.mean(y):
                            trend_direction = "上升"
                        elif slope < -0.01 * np.mean(y):
                            trend_direction = "下降"
                        else:
                            trend_direction = "平稳"
                        
                        # 计算增长率
                        if len(y) > 1 and y[0] != 0:
                            growth_rate = (y[-1] - y[0]) / abs(y[0]) * 100
                        else:
                            growth_rate = 0
                        
                        # 找出峰值和谷值
                        peak_idx = np.argmax(y)
                        valley_idx = np.argmin(y)
                        
                        peak_point = {
                            'date': str(daily_data.iloc[peak_idx]['日期']),
                            'value': float(y[peak_idx])
                        }
                        
                        valley_point = {
                            'date': str(daily_data.iloc[valley_idx]['日期']),
                            'value': float(y[valley_idx])
                        }
                        
                        analysis = TrendAnalysis(
                            time_field=time_col,
                            metric_name=metric_col,
                            data=daily_data,
                            trend_direction=trend_direction,
                            growth_rate=growth_rate,
                            peak_point=peak_point,
                            valley_point=valley_point
                        )
                        analyses.append(analysis)
                        
                    except Exception as e:
                        continue
                        
            except Exception as e:
                continue
        
        self.trend_analyses = analyses
        return analyses
    
    def _extract_key_insights(self) -> List[KeyInsight]:
        """
        基于业务指标分析关键发现
        识别异常指标，识别业务机会
        每个洞察都关联具体表格数据
        """
        insights = []
        
        if self.llm_client is None:
            # 如果没有LLM客户端，使用规则提取洞察
            return self._extract_insights_by_rules()
        
        try:
            # 准备数据摘要
            metrics_summary = self._get_metrics_summary()
            dimensions_summary = self._get_dimensions_summary()
            trends_summary = self._get_trends_summary()
            
            prompt = f"""基于以下业务数据分析结果，提取3-5个关键业务洞察。

## 业务指标摘要
{metrics_summary}

## 维度分析摘要
{dimensions_summary}

## 趋势分析摘要
{trends_summary}

请分析数据，提取以下类型的洞察：
1. 异常指标（与正常水平显著偏离的指标）
2. 业务机会（表现优异的维度或增长趋势）
3. 潜在风险（下降的趋势或表现差的维度）
4. 数据模式（值得关注的分布特征）

请以JSON格式返回，格式如下：
{{
    "insights": [
        {{
            "title": "洞察标题",
            "description": "详细描述",
            "insight_type": "anomaly/opportunity/pattern/risk",
            "supporting_data": {{"相关数据"}},
            "confidence": "高/中/低",
            "priority": "高/中/低"
        }}
    ]
}}
"""
            
            response = self.llm_client.analyze_json(
                prompt=prompt,
                system_message="你是一位专业的业务分析师，擅长从数据中发现业务洞察。",
                temperature=0.7,
                max_tokens=2000
            )
            
            if 'insights' in response:
                for item in response['insights']:
                    insight = KeyInsight(
                        title=item.get('title', ''),
                        description=item.get('description', ''),
                        insight_type=item.get('insight_type', 'pattern'),
                        supporting_data=item.get('supporting_data', {}),
                        confidence=item.get('confidence', '中'),
                        priority=item.get('priority', '中')
                    )
                    insights.append(insight)
            
        except Exception as e:
            # 如果LLM调用失败，使用规则提取
            insights = self._extract_insights_by_rules()
        
        self.key_insights = insights
        return insights
    
    def _extract_insights_by_rules(self) -> List[KeyInsight]:
        """基于规则提取洞察"""
        insights = []
        
        # 1. 分析维度分析中的Top/Bottom
        for analysis in self.dimension_analyses[:3]:
            if analysis.top_values:
                top = analysis.top_values[0]
                dim_name = analysis.dimension_name
                metric_name = analysis.metric_name
                
                insight = KeyInsight(
                    title=f"{dim_name}维度表现分析",
                    description=f"{top.get(dim_name, '未知')}在{metric_name}指标上表现最优，"
                               f"数值为{top.get('汇总', 0):.2f}。",
                    insight_type="opportunity" if top.get('汇总', 0) > 0 else "pattern",
                    supporting_data={
                        'dimension': dim_name,
                        'metric': metric_name,
                        'top_value': top
                    },
                    confidence="高",
                    priority="中"
                )
                insights.append(insight)
        
        # 2. 分析趋势
        for trend in self.trend_analyses[:2]:
            if trend.trend_direction == "上升":
                insight = KeyInsight(
                    title=f"{trend.metric_name}呈上升趋势",
                    description=f"{trend.metric_name}整体呈上升趋势，增长率为{trend.growth_rate:.2f}%。"
                               f"峰值出现在{trend.peak_point['date']}，数值为{trend.peak_point['value']:.2f}。",
                    insight_type="opportunity",
                    supporting_data={
                        'metric': trend.metric_name,
                        'growth_rate': trend.growth_rate,
                        'peak': trend.peak_point
                    },
                    confidence="高",
                    priority="高"
                )
            elif trend.trend_direction == "下降":
                insight = KeyInsight(
                    title=f"{trend.metric_name}呈下降趋势",
                    description=f"{trend.metric_name}整体呈下降趋势，增长率为{trend.growth_rate:.2f}%。"
                               f"需要关注业务风险。",
                    insight_type="risk",
                    supporting_data={
                        'metric': trend.metric_name,
                        'growth_rate': trend.growth_rate,
                        'valley': trend.valley_point
                    },
                    confidence="高",
                    priority="高"
                )
            else:
                insight = KeyInsight(
                    title=f"{trend.metric_name}保持稳定",
                    description=f"{trend.metric_name}整体保持稳定，波动较小。",
                    insight_type="pattern",
                    supporting_data={
                        'metric': trend.metric_name,
                        'growth_rate': trend.growth_rate
                    },
                    confidence="中",
                    priority="低"
                )
            insights.append(insight)
        
        # 3. 分析业务指标异常
        for metric in self.business_metrics[:3]:
            if metric.aggregation_type == 'std' and metric.value > 0:
                # 高波动性
                insight = KeyInsight(
                    title=f"{metric.name}波动较大",
                    description=f"{metric.name}的标准差为{metric.value:.2f}，表明数据波动较大，"
                               f"需要关注稳定性。",
                    insight_type="anomaly",
                    supporting_data={
                        'metric_name': metric.name,
                        'std_value': metric.value
                    },
                    confidence="中",
                    priority="中"
                )
                insights.append(insight)
                break
        
        return insights
    
    def _get_metrics_summary(self) -> str:
        """获取指标摘要"""
        lines = []
        for metric in self.business_metrics[:10]:
            value_str = metric.format_str.format(metric.value)
            lines.append(f"- {metric.name}: {value_str} {metric.unit}")
        return "\n".join(lines)
    
    def _get_dimensions_summary(self) -> str:
        """获取维度分析摘要"""
        lines = []
        for analysis in self.dimension_analyses[:5]:
            lines.append(f"维度: {analysis.dimension_name}, 指标: {analysis.metric_name}")
            if analysis.top_values:
                top = analysis.top_values[0]
                lines.append(f"  Top1: {top.get(analysis.dimension_name, 'N/A')} = {top.get('汇总', 0):.2f}")
        return "\n".join(lines)
    
    def _get_trends_summary(self) -> str:
        """获取趋势分析摘要"""
        lines = []
        for trend in self.trend_analyses[:3]:
            lines.append(f"指标: {trend.metric_name}, 时间字段: {trend.time_field}")
            lines.append(f"  趋势: {trend.trend_direction}, 增长率: {trend.growth_rate:.2f}%")
        return "\n".join(lines)
    
    def _generate_recommendations(self) -> List[Recommendation]:
        """
        基于关键洞察生成业务建议
        包括问题改进建议、机会利用建议、优化建议
        """
        recommendations = []
        
        if self.llm_client is None:
            return self._generate_recommendations_by_rules()
        
        try:
            # 准备洞察摘要
            insights_summary = []
            for insight in self.key_insights:
                insights_summary.append({
                    'title': insight.title,
                    'description': insight.description,
                    'type': insight.insight_type,
                    'priority': insight.priority
                })
            
            prompt = f"""基于以下业务洞察，生成3-5条具体的业务建议。

## 关键洞察
{json.dumps(insights_summary, ensure_ascii=False, indent=2)}

请针对每个洞察生成可执行的业务建议，建议类型包括：
1. improvement - 问题改进建议
2. opportunity - 机会利用建议
3. optimization - 优化建议

请以JSON格式返回，格式如下：
{{
    "recommendations": [
        {{
            "title": "建议标题",
            "description": "详细描述",
            "recommendation_type": "improvement/opportunity/optimization",
            "expected_impact": "预期效果",
            "implementation_difficulty": "高/中/低",
            "supporting_insights": ["关联的洞察标题"]
        }}
    ]
}}
"""
            
            response = self.llm_client.analyze_json(
                prompt=prompt,
                system_message="你是一位资深的业务顾问，擅长基于数据洞察提供可执行的业务建议。",
                temperature=0.7,
                max_tokens=2000
            )
            
            if 'recommendations' in response:
                for item in response['recommendations']:
                    rec = Recommendation(
                        title=item.get('title', ''),
                        description=item.get('description', ''),
                        recommendation_type=item.get('recommendation_type', 'improvement'),
                        expected_impact=item.get('expected_impact', ''),
                        implementation_difficulty=item.get('implementation_difficulty', '中'),
                        supporting_insights=item.get('supporting_insights', [])
                    )
                    recommendations.append(rec)
            
        except Exception as e:
            recommendations = self._generate_recommendations_by_rules()
        
        self.recommendations = recommendations
        return recommendations
    
    def _generate_recommendations_by_rules(self) -> List[Recommendation]:
        """基于规则生成建议"""
        recommendations = []
        
        for insight in self.key_insights:
            if insight.insight_type == "opportunity":
                rec = Recommendation(
                    title=f"把握{insight.title}机会",
                    description=f"基于洞察'{insight.title}'，建议加大投入，扩大优势。"
                               f"具体措施：分析成功因素，复制最佳实践到其他区域。",
                    recommendation_type="opportunity",
                    expected_impact="提升整体业绩10-20%",
                    implementation_difficulty="中",
                    supporting_insights=[insight.title]
                )
            elif insight.insight_type == "risk":
                rec = Recommendation(
                    title=f"应对{insight.title}风险",
                    description=f"针对'{insight.title}'，建议立即采取改进措施。"
                               f"具体措施：深入分析原因，制定改进计划，加强监控。",
                    recommendation_type="improvement",
                    expected_impact="遏制下滑趋势，恢复增长",
                    implementation_difficulty="高",
                    supporting_insights=[insight.title]
                )
            elif insight.insight_type == "anomaly":
                rec = Recommendation(
                    title=f"优化{insight.title}",
                    description=f"针对'{insight.title}'，建议优化流程，提升稳定性。"
                               f"具体措施：建立监控机制，识别异常原因，制定预防措施。",
                    recommendation_type="optimization",
                    expected_impact="提升稳定性，降低风险",
                    implementation_difficulty="中",
                    supporting_insights=[insight.title]
                )
            else:
                rec = Recommendation(
                    title=f"关注{insight.title}",
                    description=f"建议持续关注'{insight.title}'的发展变化，"
                               f"定期评估影响，适时调整策略。",
                    recommendation_type="optimization",
                    expected_impact="保持业务稳定",
                    implementation_difficulty="低",
                    supporting_insights=[insight.title]
                )
            recommendations.append(rec)
        
        return recommendations
    
    def _format_metric_table(self, metrics: List[BusinessMetric]) -> Dict[str, Any]:
        """
        格式化业务指标表
        表头：指标名称 | 指标值 | 单位 | 说明
        """
        headers = ["指标名称", "指标值", "单位", "说明"]
        rows = []
        
        for metric in metrics:
            value_str = metric.format_str.format(metric.value)
            rows.append([
                metric.name,
                value_str,
                metric.unit,
                metric.description
            ])
        
        return {
            'headers': headers,
            'rows': rows,
            'title': '核心业务指标'
        }
    
    def _format_dimension_table(self, analysis: DimensionAnalysis, top_n: int = 10) -> Dict[str, Any]:
        """
        格式化维度分析表
        确保数据完整无遗漏
        """
        headers = [analysis.dimension_name, '汇总', '平均', '计数']
        rows = []
        
        # 限制显示行数
        display_data = analysis.data.head(top_n)
        
        for _, row in display_data.iterrows():
            rows.append([
                str(row[analysis.dimension_name]),
                f"{row['汇总']:.2f}",
                f"{row['平均']:.2f}",
                f"{int(row['计数'])}"
            ])
        
        return {
            'headers': headers,
            'rows': rows,
            'title': f"按{analysis.dimension_name}分析 - {analysis.metric_name}"
        }
    
    def _format_trend_table(self, trend: TrendAnalysis, top_n: int = 10) -> Dict[str, Any]:
        """
        格式化趋势分析表
        """
        headers = ['日期', '数值']
        rows = []
        
        # 限制显示行数
        display_data = trend.data.head(top_n)
        
        for _, row in display_data.iterrows():
            rows.append([
                str(row['日期']),
                f"{row['数值']:.2f}"
            ])
        
        return {
            'headers': headers,
            'rows': rows,
            'title': f"{trend.metric_name}趋势分析",
            'meta': {
                'trend_direction': trend.trend_direction,
                'growth_rate': trend.growth_rate,
                'peak_point': trend.peak_point,
                'valley_point': trend.valley_point
            }
        }
    
    def _format_table_markdown(self, table_data: Dict[str, Any]) -> str:
        """生成Markdown表格语法"""
        lines = []
        
        # 表格标题
        if 'title' in table_data:
            lines.append(f"**{table_data['title']}**")
            lines.append("")
        
        headers = table_data['headers']
        rows = table_data['rows']
        
        if not headers or not rows:
            return ""
        
        # 表头
        header_line = "| " + " | ".join(headers) + " |"
        lines.append(header_line)
        
        # 分隔符
        separator = "|" + "|".join(["---" for _ in headers]) + "|"
        lines.append(separator)
        
        # 数据行
        for row in rows:
            row_line = "| " + " | ".join([str(cell) for cell in row]) + " |"
            lines.append(row_line)
        
        lines.append("")
        
        # 添加元信息
        if 'meta' in table_data:
            meta = table_data['meta']
            if 'trend_direction' in meta:
                lines.append(f"*趋势方向: {meta['trend_direction']} | 增长率: {meta['growth_rate']:.2f}%*")
                if meta.get('peak_point'):
                    lines.append(f"*峰值: {meta['peak_point']['date']} ({meta['peak_point']['value']:.2f})*")
                if meta.get('valley_point'):
                    lines.append(f"*谷值: {meta['valley_point']['date']} ({meta['valley_point']['value']:.2f})*")
            lines.append("")
        
        return "\n".join(lines)
    
    def _format_table_word(self, doc: Document, table_data: Dict[str, Any]):
        """创建Word表格对象，设置专业表格样式"""
        from docx.shared import RGBColor
        
        # 表格标题
        if 'title' in table_data:
            p = doc.add_paragraph()
            run = p.add_run(table_data['title'])
            self._set_chinese_font(run, font_size=11, bold=True)
        
        headers = table_data['headers']
        rows = table_data['rows']
        
        if not headers or not rows:
            return
        
        # 创建表格
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        table.autofit = False
        table.allow_autofit = False
        
        # 表头
        hdr_cells = table.rows[0].cells
        for j, header in enumerate(headers):
            hdr_cells[j].text = str(header)
            hdr_cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for paragraph in hdr_cells[j].paragraphs:
                for run in paragraph.runs:
                    self._set_chinese_font(run, bold=True, font_size=10)
                    run.font.color.rgb = RGBColor(255, 255, 255)
            # 深蓝色背景
            shading_elm = parse_xml(r'<w:shd {} w:fill="1F4E79"/>'.format(nsdecls('w')))
            hdr_cells[j]._tc.get_or_add_tcPr().append(shading_elm)
        
        # 数据行
        for row_idx, row in enumerate(rows):
            row_cells = table.add_row().cells
            for j, cell in enumerate(row):
                if j < len(row_cells):
                    row_cells[j].text = str(cell)
                    row_cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for paragraph in row_cells[j].paragraphs:
                        for run in paragraph.runs:
                            self._set_chinese_font(run, font_size=9)
                    # 交替行背景色
                    if row_idx % 2 == 1:
                        shading_elm = parse_xml(r'<w:shd {} w:fill="F2F2F2"/>'.format(nsdecls('w')))
                        row_cells[j]._tc.get_or_add_tcPr().append(shading_elm)
        
        doc.add_paragraph()
        
        # 添加元信息
        if 'meta' in table_data:
            meta = table_data['meta']
            if 'trend_direction' in meta:
                p = doc.add_paragraph()
                info = f"趋势方向: {meta['trend_direction']} | 增长率: {meta['growth_rate']:.2f}%"
                run = p.add_run(info)
                self._set_chinese_font(run, font_size=9, italic=True)
                
                if meta.get('peak_point'):
                    p = doc.add_paragraph()
                    info = f"峰值: {meta['peak_point']['date']} ({meta['peak_point']['value']:.2f})"
                    run = p.add_run(info)
                    self._set_chinese_font(run, font_size=9, italic=True)
                
                if meta.get('valley_point'):
                    p = doc.add_paragraph()
                    info = f"谷值: {meta['valley_point']['date']} ({meta['valley_point']['value']:.2f})"
                    run = p.add_run(info)
                    self._set_chinese_font(run, font_size=9, italic=True)
            
            doc.add_paragraph()
    
    def _set_chinese_font(self, run, font_name='Microsoft YaHei', font_size=10.5, bold=False, italic=False):
        """设置中文字体"""
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.bold = bold
        run.font.italic = italic
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    
    def generate_markdown_report(self, 
                                  title: str = "业务分析报告",
                                  metrics: Optional[List[BusinessMetric]] = None,
                                  dimensions: Optional[List[DimensionAnalysis]] = None,
                                  trends: Optional[List[TrendAnalysis]] = None,
                                  insights: Optional[List[KeyInsight]] = None,
                                  recommendations: Optional[List[Recommendation]] = None) -> str:
        """
        动态生成章节结构
        生成核心业务指标章节，动态生成各维度分析章节，生成趋势分析章节，生成关键洞察章节，生成业务建议章节
        每个章节都有规范表格支撑
        """
        # 使用传入的数据，如果没有则使用实例中的数据
        metrics = metrics or self.business_metrics
        dimensions = dimensions or self.dimension_analyses
        trends = trends or self.trend_analyses
        insights = insights or self.key_insights
        recommendations = recommendations or self.recommendations
        
        # 开始生成报告
        report_lines = []
        
        # 报告标题
        report_lines.append(f"# {title}")
        report_lines.append("")
        report_lines.append(f"**生成时间**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        report_lines.append("")
        report_lines.append("---")
        report_lines.append("")
        
        # 章节1：核心业务指标
        report_lines.append("## 一、核心业务指标")
        report_lines.append("")
        report_lines.append("本章节展示了数据中的核心业务指标，包括汇总、平均、极值等关键统计数据。")
        report_lines.append("")
        
        if metrics:
            metric_table = self._format_metric_table(metrics)
            report_lines.append(self._format_table_markdown(metric_table))
        else:
            report_lines.append("*未识别到数值型业务指标*")
            report_lines.append("")
        
        report_lines.append("---")
        report_lines.append("")
        
        # 章节2：维度分析（动态生成）
        if dimensions:
            report_lines.append("## 二、维度分析")
            report_lines.append("")
            report_lines.append("本章节按不同维度对业务指标进行分析，帮助识别各维度的表现差异。")
            report_lines.append("")
            
            for i, analysis in enumerate(dimensions[:5], 1):  # 限制显示前5个维度分析
                report_lines.append(f"### 2.{i} {analysis.dimension_name}维度分析")
                report_lines.append("")
                
                dim_table = self._format_dimension_table(analysis, top_n=10)
                report_lines.append(self._format_table_markdown(dim_table))
                
                # 添加分析说明
                if analysis.top_values:
                    top = analysis.top_values[0]
                    report_lines.append(f"**分析说明**: {top.get(analysis.dimension_name, '未知')}在该维度表现最优，"
                                       f"汇总值为{top.get('汇总', 0):.2f}。")
                    report_lines.append("")
                
                report_lines.append("")
            
            report_lines.append("---")
            report_lines.append("")
        
        # 章节3：趋势分析（动态生成）
        if trends:
            report_lines.append("## 三、趋势分析")
            report_lines.append("")
            report_lines.append("本章节展示业务指标的时间趋势变化，帮助识别增长趋势和波动规律。")
            report_lines.append("")
            
            for i, trend in enumerate(trends[:3], 1):  # 限制显示前3个趋势分析
                report_lines.append(f"### 3.{i} {trend.metric_name}趋势")
                report_lines.append("")
                
                trend_table = self._format_trend_table(trend, top_n=10)
                report_lines.append(self._format_table_markdown(trend_table))
                
                # 添加趋势说明
                report_lines.append(f"**趋势说明**: {trend.metric_name}整体呈{trend.trend_direction}趋势，"
                                   f"增长率为{trend.growth_rate:.2f}%。")
                report_lines.append("")
                
                if trend.peak_point:
                    report_lines.append(f"- 峰值出现在{trend.peak_point['date']}，数值为{trend.peak_point['value']:.2f}")
                if trend.valley_point:
                    report_lines.append(f"- 谷值出现在{trend.valley_point['date']}，数值为{trend.valley_point['value']:.2f}")
                report_lines.append("")
            
            report_lines.append("---")
            report_lines.append("")
        
        # 章节4：关键洞察
        report_lines.append("## 四、关键洞察")
        report_lines.append("")
        
        if insights:
            for i, insight in enumerate(insights, 1):
                type_names = {
                    'anomaly': '异常发现',
                    'opportunity': '业务机会',
                    'pattern': '数据模式',
                    'risk': '潜在风险'
                }
                type_name = type_names.get(insight.insight_type, insight.insight_type)
                
                report_lines.append(f"### 4.{i} {insight.title}")
                report_lines.append("")
                report_lines.append(f"**洞察类型**: {type_name}")
                report_lines.append("")
                report_lines.append(insight.description)
                report_lines.append("")
                
                # 支撑数据
                if insight.supporting_data:
                    report_lines.append("**支撑数据**:")
                    for key, value in insight.supporting_data.items():
                        if isinstance(value, dict):
                            report_lines.append(f"- {key}: {json.dumps(value, ensure_ascii=False)}")
                        else:
                            report_lines.append(f"- {key}: {value}")
                    report_lines.append("")
                
                report_lines.append(f"*置信度: {insight.confidence} | 优先级: {insight.priority}*")
                report_lines.append("")
        else:
            report_lines.append("*暂无关键洞察*")
            report_lines.append("")
        
        report_lines.append("---")
        report_lines.append("")
        
        # 章节5：业务建议
        report_lines.append("## 五、业务建议")
        report_lines.append("")
        
        if recommendations:
            for i, rec in enumerate(recommendations, 1):
                type_names = {
                    'improvement': '改进建议',
                    'opportunity': '机会建议',
                    'optimization': '优化建议'
                }
                type_name = type_names.get(rec.recommendation_type, rec.recommendation_type)
                
                report_lines.append(f"### 5.{i} {rec.title}")
                report_lines.append("")
                report_lines.append(f"**建议类型**: {type_name}")
                report_lines.append("")
                report_lines.append(rec.description)
                report_lines.append("")
                
                if rec.expected_impact:
                    report_lines.append(f"**预期效果**: {rec.expected_impact}")
                    report_lines.append("")
                
                report_lines.append(f"*实施难度: {rec.implementation_difficulty}*")
                report_lines.append("")
        else:
            report_lines.append("*暂无业务建议*")
            report_lines.append("")
        
        report_lines.append("---")
        report_lines.append("")
        
        # 报告结尾
        report_lines.append(f"*报告生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}*")
        
        return "\n".join(report_lines)
    
    def generate_word_report(self,
                              title: str = "业务分析报告",
                              metrics: Optional[List[BusinessMetric]] = None,
                              dimensions: Optional[List[DimensionAnalysis]] = None,
                              trends: Optional[List[TrendAnalysis]] = None,
                              insights: Optional[List[KeyInsight]] = None,
                              recommendations: Optional[List[Recommendation]] = None) -> io.BytesIO:
        """
        动态生成章节结构
        添加专业表格样式
        返回Word文档字节流
        """
        # 使用传入的数据，如果没有则使用实例中的数据
        metrics = metrics or self.business_metrics
        dimensions = dimensions or self.dimension_analyses
        trends = trends or self.trend_analyses
        insights = insights or self.key_insights
        recommendations = recommendations or self.recommendations
        
        # 创建Word文档
        doc = Document()
        
        # 设置默认字体
        style = doc.styles['Normal']
        style.font.name = 'Microsoft YaHei'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        
        # 标题
        title_heading = doc.add_heading(title, 0)
        for run in title_heading.runs:
            self._set_chinese_font(run, font_size=18, bold=True)
        title_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 元信息
        p = doc.add_paragraph()
        run = p.add_run(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        self._set_chinese_font(run)
        
        doc.add_paragraph()
        
        # 章节1：核心业务指标
        heading1 = doc.add_heading('一、核心业务指标', level=1)
        for run in heading1.runs:
            self._set_chinese_font(run, font_size=14, bold=True)
        
        p = doc.add_paragraph()
        run = p.add_run("本章节展示了数据中的核心业务指标，包括汇总、平均、极值等关键统计数据。")
        self._set_chinese_font(run, font_size=10)
        doc.add_paragraph()
        
        if metrics:
            metric_table = self._format_metric_table(metrics)
            self._format_table_word(doc, metric_table)
        else:
            p = doc.add_paragraph()
            run = p.add_run("未识别到数值型业务指标")
            self._set_chinese_font(run, font_size=10, italic=True)
        
        doc.add_page_break()
        
        # 章节2：维度分析（动态生成）
        if dimensions:
            heading2 = doc.add_heading('二、维度分析', level=1)
            for run in heading2.runs:
                self._set_chinese_font(run, font_size=14, bold=True)
            
            p = doc.add_paragraph()
            run = p.add_run("本章节按不同维度对业务指标进行分析，帮助识别各维度的表现差异。")
            self._set_chinese_font(run, font_size=10)
            doc.add_paragraph()
            
            for i, analysis in enumerate(dimensions[:5], 1):
                sub_heading = doc.add_heading(f"2.{i} {analysis.dimension_name}维度分析", level=2)
                for run in sub_heading.runs:
                    self._set_chinese_font(run, font_size=12, bold=True)
                
                dim_table = self._format_dimension_table(analysis, top_n=10)
                self._format_table_word(doc, dim_table)
                
                # 添加分析说明
                if analysis.top_values:
                    top = analysis.top_values[0]
                    p = doc.add_paragraph()
                    desc = f"分析说明: {top.get(analysis.dimension_name, '未知')}在该维度表现最优，汇总值为{top.get('汇总', 0):.2f}。"
                    run = p.add_run(desc)
                    self._set_chinese_font(run, font_size=9, italic=True)
                
                doc.add_paragraph()
            
            doc.add_page_break()
        
        # 章节3：趋势分析（动态生成）
        if self.trend_analyses:
            heading3 = doc.add_heading('三、趋势分析', level=1)
            for run in heading3.runs:
                self._set_chinese_font(run, font_size=14, bold=True)
            
            p = doc.add_paragraph()
            run = p.add_run("本章节展示业务指标的时间趋势变化，帮助识别增长趋势和波动规律。")
            self._set_chinese_font(run, font_size=10)
            doc.add_paragraph()
            
            for i, trend in enumerate(self.trend_analyses[:3], 1):
                sub_heading = doc.add_heading(f"3.{i} {trend.metric_name}趋势", level=2)
                for run in sub_heading.runs:
                    self._set_chinese_font(run, font_size=12, bold=True)
                
                trend_table = self._format_trend_table(trend, top_n=10)
                self._format_table_word(doc, trend_table)
                
                # 添加趋势说明
                p = doc.add_paragraph()
                desc = f"趋势说明: {trend.metric_name}整体呈{trend.trend_direction}趋势，增长率为{trend.growth_rate:.2f}%。"
                run = p.add_run(desc)
                self._set_chinese_font(run, font_size=9, italic=True)
                
                if trend.peak_point:
                    p = doc.add_paragraph()
                    info = f"峰值出现在{trend.peak_point['date']}，数值为{trend.peak_point['value']:.2f}"
                    run = p.add_run(info)
                    self._set_chinese_font(run, font_size=9)
                
                if trend.valley_point:
                    p = doc.add_paragraph()
                    info = f"谷值出现在{trend.valley_point['date']}，数值为{trend.valley_point['value']:.2f}"
                    run = p.add_run(info)
                    self._set_chinese_font(run, font_size=9)
                
                doc.add_paragraph()
            
            doc.add_page_break()
        
        # 章节4：关键洞察
        heading4 = doc.add_heading('四、关键洞察', level=1)
        for run in heading4.runs:
            self._set_chinese_font(run, font_size=14, bold=True)
        
        if insights:
            for i, insight in enumerate(insights, 1):
                type_names = {
                    'anomaly': '异常发现',
                    'opportunity': '业务机会',
                    'pattern': '数据模式',
                    'risk': '潜在风险'
                }
                type_name = type_names.get(insight.insight_type, insight.insight_type)
                
                sub_heading = doc.add_heading(f"4.{i} {insight.title}", level=2)
                for run in sub_heading.runs:
                    self._set_chinese_font(run, font_size=12, bold=True)
                
                p = doc.add_paragraph()
                run = p.add_run(f"洞察类型: {type_name}")
                self._set_chinese_font(run, font_size=10, bold=True)
                
                p = doc.add_paragraph(insight.description)
                for run in p.runs:
                    self._set_chinese_font(run, font_size=10)
                
                # 支撑数据
                if insight.supporting_data:
                    p = doc.add_paragraph()
                    run = p.add_run("支撑数据:")
                    self._set_chinese_font(run, font_size=10, bold=True)
                    
                    for key, value in insight.supporting_data.items():
                        p = doc.add_paragraph()
                        if isinstance(value, dict):
                            info = f"- {key}: {json.dumps(value, ensure_ascii=False)}"
                        else:
                            info = f"- {key}: {value}"
                        run = p.add_run(info)
                        self._set_chinese_font(run, font_size=9)
                        p.style = 'List Bullet'
                
                p = doc.add_paragraph()
                run = p.add_run(f"置信度: {insight.confidence} | 优先级: {insight.priority}")
                self._set_chinese_font(run, font_size=9)
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                
                doc.add_paragraph()
        else:
            p = doc.add_paragraph()
            run = p.add_run("暂无关键洞察")
            self._set_chinese_font(run, font_size=10, italic=True)
        
        doc.add_page_break()
        
        # 章节5：业务建议
        heading5 = doc.add_heading('五、业务建议', level=1)
        for run in heading5.runs:
            self._set_chinese_font(run, font_size=14, bold=True)
        
        if recommendations:
            for i, rec in enumerate(recommendations, 1):
                type_names = {
                    'improvement': '改进建议',
                    'opportunity': '机会建议',
                    'optimization': '优化建议'
                }
                type_name = type_names.get(rec.recommendation_type, rec.recommendation_type)
                
                sub_heading = doc.add_heading(f"5.{i} {rec.title}", level=2)
                for run in sub_heading.runs:
                    self._set_chinese_font(run, font_size=12, bold=True)
                
                p = doc.add_paragraph()
                run = p.add_run(f"建议类型: {type_name}")
                self._set_chinese_font(run, font_size=10, bold=True)
                
                p = doc.add_paragraph(rec.description)
                for run in p.runs:
                    self._set_chinese_font(run, font_size=10)
                
                if rec.expected_impact:
                    p = doc.add_paragraph()
                    run = p.add_run(f"预期效果: {rec.expected_impact}")
                    self._set_chinese_font(run, font_size=10)
                
                p = doc.add_paragraph()
                run = p.add_run(f"实施难度: {rec.implementation_difficulty}")
                self._set_chinese_font(run, font_size=9)
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                
                doc.add_paragraph()
        else:
            p = doc.add_paragraph()
            run = p.add_run("暂无业务建议")
            self._set_chinese_font(run, font_size=10, italic=True)
        
        # 页脚
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run(f"报告生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        self._set_chinese_font(run, font_size=9)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # 保存到内存
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        return doc_io


def generate_business_semantic_report(df: pd.DataFrame, 
                                       sheet_name: str = "数据",
                                       title: str = "业务分析报告",
                                       llm_client: Optional[LLMClient] = None) -> Tuple[str, io.BytesIO]:
    """
    生成业务语义分析报告的便捷函数
    
    Args:
        df: 数据DataFrame
        sheet_name: 数据表名称
        title: 报告标题
        llm_client: 可选的LLM客户端，用于生成洞察和建议
    
    Returns:
        (Markdown字符串, Word文档字节流)
    """
    generator = BusinessSemanticReportGenerator(df=df, llm_client=llm_client)
    
    # 执行完整的分析流程
    metrics = generator.analyze()
    dimensions = generator.analyze_dimensions()
    trends = generator.analyze_trends()
    insights = generator.extract_insights()
    recommendations = generator.generate_recommendations()
    
    # 生成Markdown报告
    md_report = generator.generate_markdown_report(
        title=title,
        metrics=metrics,
        dimensions=dimensions,
        trends=trends,
        insights=insights,
        recommendations=recommendations
    )
    
    # 生成Word报告
    word_report = generator.generate_word_report(
        title=title,
        metrics=metrics,
        dimensions=dimensions,
        trends=trends,
        insights=insights,
        recommendations=recommendations
    )
    
    return md_report, word_report
